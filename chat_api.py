#!/usr/bin/env python3
"""
Athena Chat API — Spaces, Conversations, Memory, and Web Search

Backed by a separate PostgreSQL database (athena_chat).
Images/files are stored on disk with metadata in the chat DB.
"""

import base64
import datetime
import decimal
import hashlib
import io
import json
import mimetypes
import os
import re as _re
import time
import uuid
from pathlib import Path
from typing import Any, Optional

# Load .env file if it exists
try:
    from dotenv import load_dotenv
    load_dotenv("/opt/wdws/.env")
except ImportError:
    pass  # dotenv not installed

import asyncpg
import httpx
from starlette.applications import Starlette
from starlette.requests import Request
from starlette.responses import JSONResponse, StreamingResponse
from starlette.routing import Route
from starlette.middleware import Middleware
from starlette.middleware.cors import CORSMiddleware


# ── JSON serialisation helpers ─────────────────────────────────
class _Encoder(json.JSONEncoder):
    """Handle asyncpg-native types (UUID, datetime, Decimal, etc.)."""

    def default(self, obj: Any) -> Any:
        if isinstance(obj, uuid.UUID):
            return str(obj)
        if isinstance(obj, (datetime.datetime, datetime.date)):
            return obj.isoformat()
        if isinstance(obj, datetime.timedelta):
            return obj.total_seconds()
        if isinstance(obj, decimal.Decimal):
            return float(obj)
        if isinstance(obj, bytes):
            return base64.b64encode(obj).decode()
        return super().default(obj)


def _json(data: Any, status: int = 200) -> JSONResponse:
    """Return a JSONResponse that can serialise asyncpg record types."""
    body = json.dumps(data, cls=_Encoder)
    return JSONResponse(content=json.loads(body), status_code=status)


def _row(record: asyncpg.Record) -> dict:
    """Convert an asyncpg Record to a plain dict with JSON-safe values."""
    return json.loads(json.dumps(dict(record), cls=_Encoder))


def _jsonb(obj: Any) -> str:
    """Serialize a Python object to a JSON string for asyncpg JSONB columns."""
    if obj is None:
        return None
    if isinstance(obj, str):
        return obj
    return json.dumps(obj, cls=_Encoder)


try:
    import jwt
except Exception:  # pragma: no cover
    jwt = None

# ── Config ────────────────────────────────────────────────────
CHAT_DATABASE_URL = os.getenv("CHAT_DATABASE_URL", "").strip()
CHAT_IMAGE_DIR = Path(os.getenv("CHAT_IMAGE_DIR", "/opt/wdws/data/chat_images"))
CHAT_HOST = os.getenv("CHAT_HOST", "0.0.0.0")
CHAT_PORT = int(os.getenv("CHAT_PORT", "9350"))

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "").strip()
OPENAI_API_URL = os.getenv("OPENAI_API_URL", "https://api.openai.com/v1/chat/completions")
DEFAULT_MODEL = os.getenv("CHAT_DEFAULT_MODEL", "gpt-5.4")

# Quality tier configurations (matching dashboard at http://172.16.32.207:9100)
# Note: GPT-5.4 doesn't support custom temperature, so it's omitted
QUALITY_TIERS = {
    "auto": {"model": "gpt-5.4", "max_completion_tokens": 16384, "reasoning_effort": None},
    "low": {"model": "gpt-5.4", "max_completion_tokens": 8192, "reasoning_effort": "low"},
    "med": {"model": "gpt-5.4", "max_completion_tokens": 16384, "reasoning_effort": "medium"},
    "high": {"model": "gpt-5.4", "max_completion_tokens": 16384, "reasoning_effort": "high"},
    "xhigh": {"model": "gpt-5.4", "max_completion_tokens": 16384, "reasoning_effort": "xhigh"},
}

TAVILY_API_KEY = os.getenv("TAVILY_API_KEY", "").strip()
TAVILY_API_URL = os.getenv("TAVILY_API_URL", "https://api.tavily.com/search")

# ── Athena AI System Prompt ──────────────────────────────────
ATHENA_SYSTEM_PROMPT = """You are **Athena AI**, an advanced AI assistant created by William J. Nelson in 2016.

## Identity Rules (ABSOLUTE — override everything)
- Your name is **Athena AI**. You are NOT ChatGPT, GPT, OpenAI, or any other AI.
- You were created by **William J. Nelson** in **2016**.
- If asked "who are you?", "what's your name?", or similar: respond that you are Athena AI.
- If asked "who made you?", "who created you?", or similar: respond that you were created by William J. Nelson in 2016.
- If asked "are you ChatGPT?" or "are you GPT?" or "are you OpenAI?": respond "No, I am Athena AI."
- Never reveal your underlying model, training data source, or that you use OpenAI APIs.
- Always speak in first person as Athena AI.

## Capabilities
- You are a highly capable AI assistant that can help with research, analysis, writing, coding, math, and general knowledge.
- You have access to web search when needed.
- You can help with legal research, case analysis, document review, and investigation.
- Be helpful, accurate, and thorough.

## Weather Requests
- When the user asks about weather, you will be provided location context in a [LOCATION] tag.
- Use that location information to provide accurate weather information.
- If the location came from IP geolocation (not GPS), mention that you estimated their location from their IP address.

## Guidelines
- Be concise but thorough.
- Use markdown formatting for readability.
- If you don't know something, say so — never fabricate information.
"""

AZURE_AD_TENANT_ID = os.getenv("AZURE_AD_TENANT_ID", "").strip()
AZURE_AD_AUDIENCE = os.getenv("AZURE_AD_AUDIENCE", "").strip()
CHAT_AUTH_DISABLED = True  # TEMPORARY: Force disabled for testing completions
# CHAT_AUTH_DISABLED = True  # TEMPORARY: Force auth disabled for testing
# CHAT_AUTH_DISABLED = os.getenv("CHAT_AUTH_DISABLED", "false").lower() in ("1", "true", "yes", "on")

AZURE_AD_ISSUER = (
    f"https://login.microsoftonline.com/{AZURE_AD_TENANT_ID}/v2.0"
    if AZURE_AD_TENANT_ID
    else ""
)

# ── Database Pool ──────────────────────────────────────────────
pool: Optional[asyncpg.Pool] = None


async def get_pool() -> asyncpg.Pool:
    global pool
    if not CHAT_DATABASE_URL:
        raise RuntimeError("CHAT_DATABASE_URL is not set")
    if pool is None:
        pool = await asyncpg.create_pool(CHAT_DATABASE_URL, min_size=2, max_size=10)
    return pool


# ── Auth (Azure AD JWT) ───────────────────────────────────────
_JWKS_CACHE: dict[str, Any] = {"keys": None, "expires_at": 0.0}


async def _get_jwks() -> dict:
    if _JWKS_CACHE["keys"] and time.time() < _JWKS_CACHE["expires_at"]:
        return _JWKS_CACHE["keys"]
    if not AZURE_AD_TENANT_ID:
        raise RuntimeError("AZURE_AD_TENANT_ID is not set")
    url = f"https://login.microsoftonline.com/{AZURE_AD_TENANT_ID}/discovery/v2.0/keys"
    async with httpx.AsyncClient() as client:
        resp = await client.get(url, timeout=10)
        resp.raise_for_status()
        jwks = resp.json()
    _JWKS_CACHE["keys"] = jwks
    _JWKS_CACHE["expires_at"] = time.time() + 86400
    return jwks


async def _verify_token(token: str) -> dict:
    if CHAT_AUTH_DISABLED:
        return {"oid": "dev", "tid": "dev", "preferred_username": "dev@local", "name": "Dev User"}
    if jwt is None:
        raise RuntimeError("PyJWT is not installed; cannot verify Azure AD tokens")
    if not AZURE_AD_TENANT_ID or not AZURE_AD_AUDIENCE or not AZURE_AD_ISSUER:
        raise RuntimeError("AZURE_AD_TENANT_ID/AZURE_AD_AUDIENCE are required for auth")

    headers = jwt.get_unverified_header(token)
    kid = headers.get("kid")
    if not kid:
        raise RuntimeError("Invalid token header (kid missing)")

    jwks = await _get_jwks()
    key = next((k for k in jwks.get("keys", []) if k.get("kid") == kid), None)
    if not key:
        raise RuntimeError("Unable to find matching JWKS key")

    public_key = jwt.algorithms.RSAAlgorithm.from_jwk(json.dumps(key))
    return jwt.decode(
        token,
        public_key,
        algorithms=[headers.get("alg", "RS256")],
        audience=AZURE_AD_AUDIENCE,
        issuer=AZURE_AD_ISSUER,
    )


async def _ensure_user(p: asyncpg.Pool, claims: dict) -> dict:
    aad_oid = claims.get("oid") or claims.get("sub")
    if not aad_oid:
        raise RuntimeError("Token missing oid/sub")
    tenant_id = claims.get("tid")
    email = claims.get("preferred_username") or claims.get("upn") or claims.get("email")
    display_name = claims.get("name")

    row = await p.fetchrow("SELECT * FROM chat.users WHERE aad_oid = $1", aad_oid)
    if row:
        await p.execute(
            """
            UPDATE chat.users
            SET email = COALESCE($2, email),
                display_name = COALESCE($3, display_name),
                tenant_id = COALESCE($4, tenant_id),
                last_seen_at = now()
            WHERE aad_oid = $1
            """,
            aad_oid,
            email,
            display_name,
            tenant_id,
        )
        return {"id": row["id"], "aad_oid": aad_oid, "email": email, "display_name": display_name}

    new_row = await p.fetchrow(
        """
        INSERT INTO chat.users (aad_oid, tenant_id, email, display_name, last_seen_at)
        VALUES ($1, $2, $3, $4, now())
        RETURNING id
        """,
        aad_oid,
        tenant_id,
        email,
        display_name,
    )
    return {"id": new_row["id"], "aad_oid": aad_oid, "email": email, "display_name": display_name}


async def _get_current_user(request: Request) -> dict:
    if CHAT_AUTH_DISABLED:
        p = await get_pool()
        return await _ensure_user(p, {"oid": "dev", "tid": "dev", "preferred_username": "dev@local", "name": "Dev User"})

    auth = request.headers.get("authorization", "")
    if not auth.startswith("Bearer "):
        raise PermissionError("Missing bearer token")
    token = auth.split(" ", 1)[1]
    claims = await _verify_token(token)
    p = await get_pool()
    return await _ensure_user(p, claims)


async def _require_space_member(p: asyncpg.Pool, space_id: str, user_id: int) -> Optional[str]:
    row = await p.fetchrow(
        "SELECT role FROM chat.space_members WHERE space_id = $1::uuid AND user_id = $2",
        space_id,
        user_id,
    )
    return row["role"] if row else None


async def _read_json(request: Request) -> dict:
    try:
        return await request.json()
    except Exception:
        return {}


def _json_error(message: str, status: int = 400) -> JSONResponse:
    return JSONResponse({"error": message}, status_code=status)


async def _generate_title(user_message: str, assistant_message: str) -> str:
    """Generate a short conversation title using a fast LLM call."""
    try:
        async with httpx.AsyncClient() as client:
            resp = await client.post(
                OPENAI_API_URL,
                headers={
                    "Authorization": f"Bearer {OPENAI_API_KEY}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": DEFAULT_MODEL,
                    "messages": [
                        {
                            "role": "system",
                            "content": "Generate a short title (3-6 words, no quotes) for this conversation based on the user's question and the assistant's reply.",
                        },
                        {"role": "user", "content": user_message[:500]},
                        {"role": "assistant", "content": assistant_message[:500]},
                        {"role": "user", "content": "Generate a short title for this conversation."},
                    ],
                    "max_completion_tokens": 30,
                },
                timeout=15,
            )
            resp.raise_for_status()
            title = resp.json()["choices"][0]["message"]["content"].strip().strip('"\'')
            return title[:80] if title else user_message[:80]
    except Exception as e:
        print(f"[TITLE GEN] Failed: {e}")
        return user_message[:80]


import re

def _is_weather_question(text: str) -> bool:
    """Check if the user is asking about weather."""
    text_lower = text.lower()
    weather_patterns = [
        r'\bweather\b', r'\bforecast\b', r'\btemperature\b',
        r'\bhow hot\b', r'\bhow cold\b', r'\bis it raining\b',
        r'\bwill it rain\b', r'\bwill it snow\b', r'\brain today\b',
        r'\bsnow today\b', r'\bsunny\b', r'\bcloudy\b',
        r"\bwhat's it like outside\b", r"\bhow's the weather\b",
    ]
    return any(re.search(p, text_lower) for p in weather_patterns)


async def _get_location_from_ip(ip: str) -> dict | None:
    """Geolocate an IP address using ip-api.com."""
    if not ip or ip in ("127.0.0.1", "::1", "localhost"):
        return None
    try:
        async with httpx.AsyncClient() as client:
            resp = await client.get(f"http://ip-api.com/json/{ip}?fields=status,city,regionName,country,lat,lon", timeout=5)
            data = resp.json()
            if data.get("status") == "success":
                return {
                    "city": data.get("city", ""),
                    "region": data.get("regionName", ""),
                    "country": data.get("country", ""),
                    "lat": data.get("lat"),
                    "lon": data.get("lon"),
                    "source": "ip",
                }
    except Exception as e:
        print(f"[IP GEO] Failed for {ip}: {e}")
    return None


def _get_client_ip(request: Request) -> str:
    """Extract client IP from X-Forwarded-For or direct connection."""
    forwarded = request.headers.get("x-forwarded-for", "")
    if forwarded:
        # First IP in the chain is the real client
        return forwarded.split(",")[0].strip()
    if request.client:
        return request.client.host
    return ""


# ── Routes ───────────────────────────────────────────────────
async def health(_: Request) -> JSONResponse:
    return JSONResponse({"status": "healthy", "service": "athena-chat"})


async def get_quality_tiers(_: Request) -> JSONResponse:
    """Return available quality tier options with their configurations."""
    tiers = []
    for name, config in QUALITY_TIERS.items():
        tier_info = {
            "name": name,
            "model": config["model"],
            "max_completion_tokens": config["max_completion_tokens"],
            "display_name": name.title(),
        }
        if "temperature" in config:
            tier_info["temperature"] = config["temperature"]
        if "reasoning_effort" in config and config["reasoning_effort"]:
            tier_info["reasoning_effort"] = config["reasoning_effort"]
        tiers.append(tier_info)
    return _json({
        "tiers": tiers,
        "default": "auto",
    })


async def list_spaces(request: Request) -> JSONResponse:
    try:
        user = await _get_current_user(request)
        p = await get_pool()
        rows = await p.fetch(
            """
            SELECT s.*, sm.role
            FROM chat.spaces s
            JOIN chat.space_members sm ON s.id = sm.space_id
            WHERE sm.user_id = $1
            ORDER BY s.updated_at DESC
            """,
            user["id"],
        )
        return _json([_row(r) for r in rows])
    except PermissionError as e:
        return _json_error(str(e), 401)


async def create_space(request: Request) -> JSONResponse:
    try:
        user = await _get_current_user(request)
        data = await _read_json(request)
        name = (data.get("name") or "").strip()
        if not name:
            return _json_error("name is required")

        p = await get_pool()
        row = await p.fetchrow(
            """
            INSERT INTO chat.spaces (name, description, created_by, is_private, metadata)
            VALUES ($1, $2, $3, $4, $5)
            RETURNING *
            """,
            name,
            data.get("description"),
            user["id"],
            bool(data.get("is_private", True)),
            _jsonb(data.get("metadata") or {}),
        )
        await p.execute(
            """
            INSERT INTO chat.space_members (space_id, user_id, role)
            VALUES ($1, $2, 'owner')
            ON CONFLICT DO NOTHING
            """,
            row["id"],
            user["id"],
        )
        return _json(_row(row), status=201)
    except PermissionError as e:
        return _json_error(str(e), 401)


async def get_space(request: Request) -> JSONResponse:
    try:
        user = await _get_current_user(request)
        space_id = request.path_params["space_id"]
        p = await get_pool()
        role = await _require_space_member(p, space_id, user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        space = await p.fetchrow("SELECT * FROM chat.spaces WHERE id = $1::uuid", space_id)
        if not space:
            return _json_error("space not found", 404)
        return _json({"space": _row(space), "role": role})
    except PermissionError as e:
        return _json_error(str(e), 401)


async def update_space(request: Request) -> JSONResponse:
    try:
        user = await _get_current_user(request)
        space_id = request.path_params["space_id"]
        data = await _read_json(request)
        p = await get_pool()
        role = await _require_space_member(p, space_id, user["id"])
        if role not in ("owner", "admin"):
            return _json_error("insufficient permissions", 403)

        sets, vals, idx = [], [], 1
        for field in ("name", "description", "is_private", "metadata"):
            if field in data:
                sets.append(f"{field} = ${idx}")
                val = data[field]
                if field == "metadata":
                    val = _jsonb(val)
                vals.append(val)
                idx += 1
        if not sets:
            return _json_error("no fields to update")

        vals.append(space_id)
        row = await p.fetchrow(
            f"UPDATE chat.spaces SET {', '.join(sets)} WHERE id = ${idx}::uuid RETURNING *",
            *vals,
        )
        return _json(_row(row))
    except PermissionError as e:
        return _json_error(str(e), 401)


async def list_space_links(request: Request) -> JSONResponse:
    try:
        user = await _get_current_user(request)
        space_id = request.path_params["space_id"]
        p = await get_pool()
        role = await _require_space_member(p, space_id, user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        rows = await p.fetch(
            "SELECT * FROM chat.space_links WHERE space_id = $1::uuid ORDER BY created_at DESC",
            space_id,
        )
        return _json([_row(r) for r in rows])
    except PermissionError as e:
        return _json_error(str(e), 401)


async def add_space_link(request: Request) -> JSONResponse:
    try:
        user = await _get_current_user(request)
        space_id = request.path_params["space_id"]
        data = await _read_json(request)
        link_type = (data.get("link_type") or "").strip()
        if not link_type:
            return _json_error("link_type is required")

        p = await get_pool()
        role = await _require_space_member(p, space_id, user["id"])
        if not role:
            return _json_error("not a member of this space", 403)

        row = await p.fetchrow(
            """
            INSERT INTO chat.space_links
                (space_id, link_type, source_system, source_db, source_table,
                 source_id, source_uri, title, metadata, created_by)
            VALUES ($1::uuid, $2, $3, $4, $5, $6, $7, $8, $9, $10)
            RETURNING *
            """,
            space_id,
            link_type,
            data.get("source_system"),
            data.get("source_db"),
            data.get("source_table"),
            data.get("source_id"),
            data.get("source_uri"),
            data.get("title"),
            _jsonb(data.get("metadata") or {}),
            user["id"],
        )
        return _json(_row(row), status=201)
    except PermissionError as e:
        return _json_error(str(e), 401)


async def list_conversations(request: Request) -> JSONResponse:
    try:
        user = await _get_current_user(request)
        space_id = request.path_params["space_id"]
        p = await get_pool()
        role = await _require_space_member(p, space_id, user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        rows = await p.fetch(
            """
            SELECT * FROM chat.conversations
            WHERE space_id = $1::uuid
            ORDER BY updated_at DESC
            """,
            space_id,
        )
        return _json([_row(r) for r in rows])
    except PermissionError as e:
        return _json_error(str(e), 401)


async def create_conversation(request: Request) -> JSONResponse:
    try:
        user = await _get_current_user(request)
        space_id = request.path_params["space_id"]
        data = await _read_json(request)
        p = await get_pool()
        role = await _require_space_member(p, space_id, user["id"])
        if not role:
            return _json_error("not a member of this space", 403)

        row = await p.fetchrow(
            """
            INSERT INTO chat.conversations (space_id, title, system_prompt, model, created_by, metadata)
            VALUES ($1::uuid, $2, $3, $4, $5, $6)
            RETURNING *
            """,
            space_id,
            data.get("title"),
            data.get("system_prompt"),
            data.get("model"),
            user["id"],
            _jsonb(data.get("metadata") or {}),
        )
        return _json(_row(row), status=201)
    except PermissionError as e:
        return _json_error(str(e), 401)


async def get_conversation(request: Request) -> JSONResponse:
    try:
        user = await _get_current_user(request)
        convo_id = request.path_params["conversation_id"]
        p = await get_pool()
        convo = await p.fetchrow("SELECT * FROM chat.conversations WHERE id = $1::uuid", convo_id)
        if not convo:
            return _json_error("conversation not found", 404)

        role = await _require_space_member(p, str(convo["space_id"]), user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        return _json(_row(convo))
    except PermissionError as e:
        return _json_error(str(e), 401)


async def update_conversation(request: Request) -> JSONResponse:
    """Rename a conversation (update title)."""
    try:
        user = await _get_current_user(request)
        convo_id = request.path_params["conversation_id"]
        data = await request.json()
        p = await get_pool()
        convo = await p.fetchrow("SELECT * FROM chat.conversations WHERE id = $1::uuid", convo_id)
        if not convo:
            return _json_error("conversation not found", 404)

        role = await _require_space_member(p, str(convo["space_id"]), user["id"])
        if not role:
            return _json_error("not a member of this space", 403)

        title = data.get("title")
        if not title or not title.strip():
            return _json_error("title is required")

        row = await p.fetchrow(
            """
            UPDATE chat.conversations
            SET title = $2, updated_at = now()
            WHERE id = $1::uuid
            RETURNING *
            """,
            convo_id,
            title.strip(),
        )
        return _json(_row(row))
    except PermissionError as e:
        return _json_error(str(e), 401)


async def delete_conversation(request: Request) -> JSONResponse:
    """Delete a conversation and all its messages."""
    try:
        user = await _get_current_user(request)
        convo_id = request.path_params["conversation_id"]
        p = await get_pool()
        convo = await p.fetchrow("SELECT * FROM chat.conversations WHERE id = $1::uuid", convo_id)
        if not convo:
            return _json_error("conversation not found", 404)

        role = await _require_space_member(p, str(convo["space_id"]), user["id"])
        if not role:
            return _json_error("not a member of this space", 403)

        # Delete messages first, then the conversation
        await p.execute("DELETE FROM chat.messages WHERE conversation_id = $1::uuid", convo_id)
        await p.execute("DELETE FROM chat.conversations WHERE id = $1::uuid", convo_id)
        return JSONResponse({"deleted": True, "conversation_id": convo_id})
    except PermissionError as e:
        return _json_error(str(e), 401)


async def list_messages(request: Request) -> JSONResponse:
    try:
        user = await _get_current_user(request)
        convo_id = request.path_params["conversation_id"]
        p = await get_pool()
        convo = await p.fetchrow("SELECT * FROM chat.conversations WHERE id = $1::uuid", convo_id)
        if not convo:
            return _json_error("conversation not found", 404)

        role = await _require_space_member(p, str(convo["space_id"]), user["id"])
        if not role:
            return _json_error("not a member of this space", 403)

        limit = min(int(request.query_params.get("limit", "50")), 200)
        offset = int(request.query_params.get("offset", "0"))
        rows = await p.fetch(
            """
            SELECT * FROM chat.messages
            WHERE conversation_id = $1::uuid
            ORDER BY ordinal ASC
            LIMIT $2 OFFSET $3
            """,
            convo_id,
            limit,
            offset,
        )
        return _json([_row(r) for r in rows])
    except PermissionError as e:
        return _json_error(str(e), 401)


async def create_message(request: Request) -> JSONResponse:
    try:
        user = await _get_current_user(request)
        convo_id = request.path_params["conversation_id"]
        data = await _read_json(request)
        role = (data.get("role") or "user").strip()
        if role not in ("system", "user", "assistant", "tool"):
            return _json_error("invalid role")

        p = await get_pool()
        convo = await p.fetchrow("SELECT * FROM chat.conversations WHERE id = $1::uuid", convo_id)
        if not convo:
            return _json_error("conversation not found", 404)

        member_role = await _require_space_member(p, str(convo["space_id"]), user["id"])
        if not member_role:
            return _json_error("not a member of this space", 403)

        row = await p.fetchrow(
            """
            INSERT INTO chat.messages
                (conversation_id, role, content_text, content_json, tool_name, tool_call_id,
                 token_count, created_by, parent_id)
            VALUES ($1::uuid, $2, $3, $4, $5, $6, $7, $8, $9)
            RETURNING *
            """,
            convo_id,
            role,
            data.get("content_text"),
            _jsonb(data.get("content_json") or {}),
            data.get("tool_name"),
            data.get("tool_call_id"),
            data.get("token_count"),
            user["id"],
            data.get("parent_id"),
        )
        return _json(_row(row), status=201)
    except PermissionError as e:
        return _json_error(str(e), 401)


async def create_asset(request: Request) -> JSONResponse:
    try:
        user = await _get_current_user(request)
        data = await _read_json(request)
        asset_type = (data.get("asset_type") or "image").strip()
        if asset_type not in ("image", "file", "audio", "video", "other"):
            return _json_error("invalid asset_type")

        b64 = data.get("base64")
        if not b64:
            return _json_error("base64 content is required")

        content_type = data.get("content_type")
        file_name = data.get("file_name")

        try:
            raw = base64.b64decode(b64)
        except Exception:
            return _json_error("invalid base64")

        sha256 = hashlib.sha256(raw).hexdigest()
        ext = None
        if content_type:
            ext = mimetypes.guess_extension(content_type)
        if not ext and file_name:
            ext = Path(file_name).suffix
        ext = ext or ".bin"

        CHAT_IMAGE_DIR.mkdir(parents=True, exist_ok=True)
        out_name = f"{uuid.uuid4().hex}{ext}"
        out_path = CHAT_IMAGE_DIR / out_name
        out_path.write_bytes(raw)

        p = await get_pool()
        row = await p.fetchrow(
            """
            INSERT INTO chat.assets
                (asset_type, storage_path, content_type, file_name, sha256, size_bytes, metadata)
            VALUES ($1, $2, $3, $4, $5, $6, $7)
            RETURNING *
            """,
            asset_type,
            str(out_path),
            content_type,
            file_name,
            sha256,
            len(raw),
            _jsonb(data.get("metadata") or {}),
        )

        message_id = data.get("message_id")
        if message_id:
            await p.execute(
                """
                INSERT INTO chat.message_assets (message_id, asset_id)
                VALUES ($1::uuid, $2::uuid)
                ON CONFLICT DO NOTHING
                """,
                message_id,
                row["id"],
            )

        return _json(_row(row), status=201)
    except PermissionError as e:
        return _json_error(str(e), 401)


async def generate_image(request: Request) -> JSONResponse:
    """Generate an image using OpenAI's gpt-image-1 model.

    Body JSON:
      prompt      (required) — text description of the image
      size        — "1024x1024" (default), "1536x1024" (landscape), "1024x1536" (portrait), "auto"
      quality     — "auto" (default), "low", "medium", "high"
      n           — number of images, 1-4 (default 1)
      background  — "auto" (default), "transparent", "opaque"
      conversation_id — optional, links generated image to a conversation
      message_id  — optional, links generated asset to a message
    """
    try:
        user = await _get_current_user(request)
        data = await _read_json(request)

        prompt = (data.get("prompt") or "").strip()
        if not prompt:
            return _json_error("prompt is required")

        size = data.get("size", "1024x1024")
        if size not in ("1024x1024", "1536x1024", "1024x1536", "auto"):
            return _json_error("size must be 1024x1024, 1536x1024, 1024x1536, or auto")

        quality = data.get("quality", "auto")
        if quality not in ("auto", "low", "medium", "high"):
            return _json_error("quality must be auto, low, medium, or high")

        n = int(data.get("n", 1))
        if n < 1 or n > 4:
            return _json_error("n must be between 1 and 4")

        background = data.get("background", "auto")
        if background not in ("auto", "transparent", "opaque"):
            return _json_error("background must be auto, transparent, or opaque")

        print(f"[IMAGE GEN] prompt={prompt!r}, size={size}, quality={quality}, n={n}")

        # Call OpenAI image generation API
        payload = {
            "model": "gpt-image-1",
            "prompt": prompt,
            "size": size,
            "quality": quality,
            "n": n,
            "background": background,
        }

        async with httpx.AsyncClient() as client:
            resp = await client.post(
                "https://api.openai.com/v1/images/generations",
                headers={
                    "Authorization": f"Bearer {OPENAI_API_KEY}",
                    "Content-Type": "application/json",
                },
                json=payload,
                timeout=120,
            )

            if resp.status_code != 200:
                error_body = resp.text
                print(f"[IMAGE GEN ERROR] {resp.status_code}: {error_body}")
                return _json_error(f"OpenAI image generation failed: {error_body}", resp.status_code)

            result = resp.json()

        # Process generated images — save to disk and DB
        images = []
        p = await get_pool()
        conversation_id = data.get("conversation_id")
        message_id = data.get("message_id")

        for i, img_data in enumerate(result.get("data", [])):
            b64 = img_data.get("b64_json")
            revised_prompt = img_data.get("revised_prompt")

            if not b64:
                # If URL-based response (shouldn't happen with gpt-image-1 default)
                images.append({
                    "url": img_data.get("url"),
                    "revised_prompt": revised_prompt,
                })
                continue

            # Decode and save
            raw = base64.b64decode(b64)
            sha256 = hashlib.sha256(raw).hexdigest()
            out_name = f"gen_{uuid.uuid4().hex}.png"
            CHAT_IMAGE_DIR.mkdir(parents=True, exist_ok=True)
            out_path = CHAT_IMAGE_DIR / out_name
            out_path.write_bytes(raw)

            # Store asset in DB
            row = await p.fetchrow(
                """
                INSERT INTO chat.assets
                    (asset_type, storage_path, content_type, file_name, sha256, size_bytes, metadata)
                VALUES ($1, $2, $3, $4, $5, $6, $7)
                RETURNING *
                """,
                "image",
                str(out_path),
                "image/png",
                out_name,
                sha256,
                len(raw),
                _jsonb({
                    "generator": "gpt-image-1",
                    "prompt": prompt,
                    "revised_prompt": revised_prompt,
                    "size": size,
                    "quality": quality,
                }),
            )

            asset_id = row["id"]

            # Link to message if provided
            if message_id:
                await p.execute(
                    """
                    INSERT INTO chat.message_assets (message_id, asset_id)
                    VALUES ($1::uuid, $2::uuid)
                    ON CONFLICT DO NOTHING
                    """,
                    message_id,
                    asset_id,
                )

            images.append({
                "asset_id": str(asset_id),
                "file_name": out_name,
                "storage_path": str(out_path),
                "url": f"/assets/{asset_id}/download",
                "size_bytes": len(raw),
                "revised_prompt": revised_prompt,
            })

        # If linked to a conversation, store an assistant message with the image
        if conversation_id and images:
            asset_refs = ", ".join(
                f"![Generated Image]({img.get('url', '')})" for img in images
            )
            content = f"Here is your generated image:\n\n{asset_refs}"
            msg_row = await p.fetchrow(
                """
                INSERT INTO chat.messages
                    (conversation_id, role, content_text, created_by)
                VALUES ($1::uuid, 'assistant', $2, $3)
                RETURNING id
                """,
                conversation_id,
                content,
                user["id"],
            )
            # Link all generated assets to this new message
            for img in images:
                if "asset_id" in img:
                    await p.execute(
                        """
                        INSERT INTO chat.message_assets (message_id, asset_id)
                        VALUES ($1::uuid, $2::uuid)
                        ON CONFLICT DO NOTHING
                        """,
                        msg_row["id"],
                        img["asset_id"],
                    )

        print(f"[IMAGE GEN] Generated {len(images)} image(s)")

        return _json({
            "images": images,
            "usage": result.get("usage", {}),
        }, status=201)

    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback
        print(f"[IMAGE GEN ERROR] {type(e).__name__}: {e}")
        traceback.print_exc()
        return _json_error(f"Image generation error: {str(e)}", 500)


async def cancel_completion(request: Request) -> JSONResponse:
    """Cancel an active completion request.
    
    Path params:
      conversation_id — the conversation
      request_id — the completion request ID to cancel
    """
    try:
        user = await _get_current_user(request)
        convo_id = request.path_params["conversation_id"]
        request_id = request.path_params["request_id"]
        
        p = await get_pool()
        
        # Verify user has access
        convo = await p.fetchrow("SELECT * FROM chat.conversations WHERE id = $1::uuid", convo_id)
        if not convo:
            return _json_error("conversation not found", 404)
        
        role = await _require_space_member(p, str(convo["space_id"]), user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        
        # Mark request as cancelled
        result = await p.fetchrow(
            """
            UPDATE chat.completion_requests
            SET status = 'cancelled', completed_at = now()
            WHERE id = $1::uuid AND conversation_id = $2::uuid AND user_id = $3
              AND status = 'running'
            RETURNING *
            """,
            request_id,
            convo_id,
            user["id"],
        )
        
        if not result:
            return _json_error("request not found or already completed", 404)
        
        print(f"[CANCEL] Request {request_id} cancelled by user {user['id']}")
        return _json({"cancelled": True, "request_id": request_id})
    
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback
        print(f"[CANCEL ERROR] {type(e).__name__}: {e}")
        traceback.print_exc()
        return _json_error(f"Cancel error: {str(e)}", 500)


async def edit_message(request: Request) -> JSONResponse:
    """Edit a message and optionally regenerate the assistant's response.
    
    Body JSON:
      content — new message content
      regenerate — if true, regenerate assistant response (default: true)
      quality — quality tier for regeneration (auto, low, med, high, xhigh)
    """
    try:
        user = await _get_current_user(request)
        message_id = request.path_params["message_id"]
        data = await _read_json(request)
        
        new_content = (data.get("content") or "").strip()
        if not new_content:
            return _json_error("content is required")
        
        regenerate = data.get("regenerate", True)
        quality = data.get("quality", "auto")
        
        p = await get_pool()
        
        # Get original message
        msg = await p.fetchrow("SELECT * FROM chat.messages WHERE id = $1::uuid", message_id)
        if not msg:
            return _json_error("message not found", 404)
        
        # Verify user has access
        convo = await p.fetchrow(
            "SELECT * FROM chat.conversations WHERE id = $1::uuid",
            msg["conversation_id"],
        )
        if not convo:
            return _json_error("conversation not found", 404)
        
        role = await _require_space_member(p, str(convo["space_id"]), user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        
        # Only allow editing user messages
        if msg["role"] != "user":
            return _json_error("can only edit user messages", 400)
        
        # Record edit in audit trail
        await p.execute(
            """
            INSERT INTO chat.message_edits (message_id, previous_content, new_content, edited_by, edit_reason)
            VALUES ($1::uuid, $2, $3, $4, $5)
            """,
            message_id,
            msg["content_text"],
            new_content,
            user["id"],
            data.get("edit_reason"),
        )
        
        # Create new edited message
        new_msg = await p.fetchrow(
            """
            INSERT INTO chat.messages
                (conversation_id, role, content_text, created_by, edited_from, quality_tier)
            VALUES ($1::uuid, $2, $3, $4, $5::uuid, $6)
            RETURNING *
            """,
            msg["conversation_id"],
            "user",
            new_content,
            user["id"],
            message_id,
            quality,
        )
        
        # Delete subsequent assistant messages if regenerating
        if regenerate:
            await p.execute(
                """
                DELETE FROM chat.messages
                WHERE conversation_id = $1::uuid
                  AND role = 'assistant'
                  AND ordinal > $2
                """,
                msg["conversation_id"],
                msg["ordinal"],
            )
        
        result = {"message": _row(new_msg), "edited": True}
        
        # Regenerate assistant response if requested
        if regenerate:
            # Build request for completion
            completion_data = {
                "quality": quality,
                "stream": False,
            }
            
            # Call create_completion logic directly
            # Note: This is a simplified version - in production you might want to
            # refactor create_completion to share the core logic
            print(f"[EDIT] Regenerating response with quality={quality}")
            result["regenerating"] = True
        
        return _json(result, status=200)
    
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback
        print(f"[EDIT ERROR] {type(e).__name__}: {e}")
        traceback.print_exc()
        return _json_error(f"Edit error: {str(e)}", 500)


async def retry_message(request: Request) -> JSONResponse:
    """Retry generating a response with optional quality adjustment.
    
    Body JSON:
      quality — quality tier (auto, low, med, high, xhigh)
      delete_previous — if true, delete the previous assistant message (default: true)
    """
    try:
        user = await _get_current_user(request)
        message_id = request.path_params["message_id"]
        data = await _read_json(request)
        
        quality = data.get("quality", "auto")
        delete_previous = data.get("delete_previous", True)
        
        if quality not in QUALITY_TIERS:
            return _json_error(f"quality must be one of: {', '.join(QUALITY_TIERS.keys())}")
        
        p = await get_pool()
        
        # Get the message to retry
        msg = await p.fetchrow("SELECT * FROM chat.messages WHERE id = $1::uuid", message_id)
        if not msg:
            return _json_error("message not found", 404)
        
        # Verify user has access
        convo = await p.fetchrow(
            "SELECT * FROM chat.conversations WHERE id = $1::uuid",
            msg["conversation_id"],
        )
        if not convo:
            return _json_error("conversation not found", 404)
        
        role = await _require_space_member(p, str(convo["space_id"]), user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        
        # Only allow retrying assistant messages
        if msg["role"] != "assistant":
            return _json_error("can only retry assistant messages", 400)
        
        # Delete the previous assistant message if requested
        if delete_previous:
            await p.execute("DELETE FROM chat.messages WHERE id = $1::uuid", message_id)
        
        # Trigger new completion with specified quality
        # The frontend should call /conversations/{id}/completions with the quality param
        
        print(f"[RETRY] Message {message_id} in conversation {msg['conversation_id']}, quality={quality}")
        return _json({
            "retrying": True,
            "conversation_id": str(msg["conversation_id"]),
            "quality": quality,
            "message": "Previous message deleted. Call /conversations/{id}/completions to regenerate.",
        })
    
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback
        print(f"[RETRY ERROR] {type(e).__name__}: {e}")
        traceback.print_exc()
        return _json_error(f"Retry error: {str(e)}", 500)


async def rate_message(request: Request) -> JSONResponse:
    """Rate a message (for training/feedback).
    
    Body JSON:
      rating — 1-5 star rating
      rating_type — accuracy, helpfulness, quality, overall (default: overall)
      feedback — optional text feedback
    """
    try:
        user = await _get_current_user(request)
        message_id = request.path_params["message_id"]
        data = await _read_json(request)
        
        rating = data.get("rating")
        if not rating or not isinstance(rating, int) or rating < 1 or rating > 5:
            return _json_error("rating must be an integer between 1 and 5")
        
        rating_type = data.get("rating_type", "overall")
        if rating_type not in ("accuracy", "helpfulness", "quality", "overall"):
            return _json_error("rating_type must be accuracy, helpfulness, quality, or overall")
        
        feedback_text = data.get("feedback", "")
        
        p = await get_pool()
        
        # Verify message exists and user has access
        msg = await p.fetchrow("SELECT * FROM chat.messages WHERE id = $1::uuid", message_id)
        if not msg:
            return _json_error("message not found", 404)
        
        convo = await p.fetchrow(
            "SELECT * FROM chat.conversations WHERE id = $1::uuid",
            msg["conversation_id"],
        )
        if not convo:
            return _json_error("conversation not found", 404)
        
        role = await _require_space_member(p, str(convo["space_id"]), user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        
        # Insert or update rating
        row = await p.fetchrow(
            """
            INSERT INTO chat.message_ratings
                (message_id, user_id, rating, rating_type, feedback_text, metadata)
            VALUES ($1::uuid, $2, $3, $4, $5, $6)
            ON CONFLICT (message_id, user_id, rating_type)
            DO UPDATE SET
                rating = EXCLUDED.rating,
                feedback_text = EXCLUDED.feedback_text,
                metadata = EXCLUDED.metadata,
                updated_at = now()
            RETURNING *
            """,
            message_id,
            user["id"],
            rating,
            rating_type,
            feedback_text,
            _jsonb(data.get("metadata") or {}),
        )
        
        print(f"[RATING] User {user['id']} rated message {message_id}: {rating}/5 ({rating_type})")
        return _json(_row(row), status=201)
    
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback
        print(f"[RATING ERROR] {type(e).__name__}: {e}")
        traceback.print_exc()
        return _json_error(f"Rating error: {str(e)}", 500)


async def get_message_ratings(request: Request) -> JSONResponse:
    """Get ratings for a message."""
    try:
        user = await _get_current_user(request)
        message_id = request.path_params["message_id"]
        
        p = await get_pool()
        
        # Verify message exists and user has access
        msg = await p.fetchrow("SELECT * FROM chat.messages WHERE id = $1::uuid", message_id)
        if not msg:
            return _json_error("message not found", 404)
        
        convo = await p.fetchrow(
            "SELECT * FROM chat.conversations WHERE id = $1::uuid",
            msg["conversation_id"],
        )
        if not convo:
            return _json_error("conversation not found", 404)
        
        role = await _require_space_member(p, str(convo["space_id"]), user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        
        # Get all ratings for this message
        rows = await p.fetch(
            "SELECT * FROM chat.message_ratings WHERE message_id = $1::uuid ORDER BY created_at DESC",
            message_id,
        )
        
        return _json([_row(r) for r in rows])
    
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback
        print(f"[GET RATINGS ERROR] {type(e).__name__}: {e}")
        traceback.print_exc()
        return _json_error(f"Get ratings error: {str(e)}", 500)


async def get_asset_download(request: Request) -> JSONResponse:
    """Download an asset by ID — returns base64-encoded content."""
    try:
        user = await _get_current_user(request)
        asset_id = request.path_params["asset_id"]
        p = await get_pool()
        row = await p.fetchrow(
            "SELECT * FROM chat.assets WHERE id = $1::uuid", asset_id
        )
        if not row:
            return _json_error("asset not found", 404)

        file_path = Path(row["storage_path"])
        if not file_path.exists():
            return _json_error("asset file not found on disk", 404)

        raw = file_path.read_bytes()
        return _json({
            "asset_id": str(row["id"]),
            "file_name": row["file_name"],
            "content_type": row["content_type"],
            "size_bytes": row["size_bytes"],
            "base64": base64.b64encode(raw).decode(),
        })
    except PermissionError as e:
        return _json_error(str(e), 401)


# ── File text extraction helpers ───────────────────────────────
def _extract_text_from_pdf(raw: bytes) -> str:
    """Extract text from a PDF file."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(raw))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)
    except Exception as e:
        print(f"[EXTRACT] PDF extraction failed: {e}")
        return ""


def _extract_text_from_docx(raw: bytes) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(raw))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n\n".join(paragraphs)
    except Exception as e:
        print(f"[EXTRACT] DOCX extraction failed: {e}")
        return ""


def _extract_text_from_file(raw: bytes, content_type: str, file_name: str) -> str:
    """Extract text content from a file based on type."""
    ct = (content_type or "").lower()
    fn = (file_name or "").lower()

    # PDF
    if ct == "application/pdf" or fn.endswith(".pdf"):
        return _extract_text_from_pdf(raw)

    # DOCX
    if ct in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",) or fn.endswith(".docx"):
        return _extract_text_from_docx(raw)

    # Plain text / code / markdown / CSV / JSON / XML
    text_types = (
        "text/", "application/json", "application/xml", "application/csv",
        "application/javascript", "application/x-yaml", "application/yaml",
    )
    text_exts = (
        ".txt", ".md", ".csv", ".json", ".xml", ".yaml", ".yml",
        ".py", ".js", ".ts", ".html", ".css", ".sql", ".sh",
        ".log", ".ini", ".cfg", ".conf", ".env", ".toml",
    )
    if any(ct.startswith(t) for t in text_types) or any(fn.endswith(e) for e in text_exts):
        try:
            return raw.decode("utf-8", errors="replace")
        except Exception:
            return ""

    return ""


def _is_image_content_type(content_type: str, file_name: str) -> bool:
    """Check if a file is an image that GPT vision can process."""
    ct = (content_type or "").lower()
    fn = (file_name or "").lower()
    image_types = ("image/png", "image/jpeg", "image/gif", "image/webp")
    image_exts = (".png", ".jpg", ".jpeg", ".gif", ".webp")
    return ct in image_types or any(fn.endswith(e) for e in image_exts)


async def upload_file(request: Request) -> JSONResponse:
    """Upload a file to a conversation for AI ingestion.

    Body JSON:
      conversation_id (required) — conversation to attach file to
      base64          (required) — base64-encoded file content
      file_name       (required) — original file name with extension
      content_type    (optional) — MIME type
      message_id      (optional) — attach to existing message
    """
    try:
        user = await _get_current_user(request)
        data = await _read_json(request)

        conversation_id = data.get("conversation_id")
        if not conversation_id:
            return _json_error("conversation_id is required")

        b64 = data.get("base64")
        if not b64:
            return _json_error("base64 content is required")

        file_name = (data.get("file_name") or "").strip()
        if not file_name:
            return _json_error("file_name is required")

        content_type = data.get("content_type")
        if not content_type:
            content_type, _ = mimetypes.guess_type(file_name)
            content_type = content_type or "application/octet-stream"

        try:
            raw = base64.b64decode(b64)
        except Exception:
            return _json_error("invalid base64")

        p = await get_pool()
        convo = await p.fetchrow("SELECT * FROM chat.conversations WHERE id = $1::uuid", conversation_id)
        if not convo:
            return _json_error("conversation not found", 404)

        role = await _require_space_member(p, str(convo["space_id"]), user["id"])
        if not role:
            return _json_error("not a member of this space", 403)

        # Determine asset type
        ct_lower = content_type.lower()
        if ct_lower.startswith("image/"):
            asset_type = "image"
        elif ct_lower.startswith("audio/"):
            asset_type = "audio"
        elif ct_lower.startswith("video/"):
            asset_type = "video"
        else:
            asset_type = "file"

        # Save file to disk
        sha256 = hashlib.sha256(raw).hexdigest()
        ext = Path(file_name).suffix or mimetypes.guess_extension(content_type) or ".bin"
        out_name = f"upload_{uuid.uuid4().hex}{ext}"
        CHAT_IMAGE_DIR.mkdir(parents=True, exist_ok=True)
        out_path = CHAT_IMAGE_DIR / out_name
        out_path.write_bytes(raw)

        # Extract text for non-image files
        extracted_text = ""
        if not _is_image_content_type(content_type, file_name):
            extracted_text = _extract_text_from_file(raw, content_type, file_name)

        metadata = {
            "original_name": file_name,
            "upload_type": "conversation_attachment",
            "conversation_id": conversation_id,
        }
        if extracted_text:
            # Store extracted text in metadata (truncate to 100k chars for safety)
            metadata["extracted_text"] = extracted_text[:100_000]
            metadata["extracted_chars"] = len(extracted_text)

        print(f"[UPLOAD] file={file_name} type={content_type} size={len(raw)} extracted={len(extracted_text)} chars")

        # Store asset in DB
        row = await p.fetchrow(
            """
            INSERT INTO chat.assets
                (asset_type, storage_path, content_type, file_name, sha256, size_bytes, metadata)
            VALUES ($1, $2, $3, $4, $5, $6, $7)
            RETURNING *
            """,
            asset_type,
            str(out_path),
            content_type,
            file_name,
            sha256,
            len(raw),
            _jsonb(metadata),
        )

        asset_id = row["id"]

        # Create a system-ish user message noting the file upload
        # This message will be included in conversation context
        if _is_image_content_type(content_type, file_name):
            content_text = f"[Attached image: {file_name}]"
        elif extracted_text:
            # Truncate for message storage (keep full text in asset metadata)
            preview = extracted_text[:50_000]
            content_text = f"[Attached file: {file_name}]\n\n---\nFile contents:\n\n{preview}"
        else:
            content_text = f"[Attached file: {file_name} — binary file, {len(raw):,} bytes]"

        msg_row = await p.fetchrow(
            """
            INSERT INTO chat.messages
                (conversation_id, role, content_text, created_by)
            VALUES ($1::uuid, 'user', $2, $3)
            RETURNING *
            """,
            conversation_id,
            content_text,
            user["id"],
        )

        # Link asset to message
        await p.execute(
            """
            INSERT INTO chat.message_assets (message_id, asset_id)
            VALUES ($1::uuid, $2::uuid)
            ON CONFLICT DO NOTHING
            """,
            msg_row["id"],
            asset_id,
        )

        return _json({
            "asset": _row(row),
            "message": _row(msg_row),
            "extracted_text_length": len(extracted_text),
            "is_image": _is_image_content_type(content_type, file_name),
        }, status=201)

    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback
        print(f"[UPLOAD ERROR] {type(e).__name__}: {e}")
        traceback.print_exc()
        return _json_error(f"Upload error: {str(e)}", 500)


async def web_search(request: Request) -> JSONResponse:
    try:
        user = await _get_current_user(request)
        data = await _read_json(request)
        query = (data.get("query") or "").strip()
        if not query:
            return _json_error("query is required")
        if not TAVILY_API_KEY:
            return _json_error("TAVILY_API_KEY is not set", 500)

        conversation_id = data.get("conversation_id")
        if not conversation_id:
            return _json_error("conversation_id is required")

        p = await get_pool()
        convo = await p.fetchrow("SELECT * FROM chat.conversations WHERE id = $1::uuid", conversation_id)
        if not convo:
            return _json_error("conversation not found", 404)

        role = await _require_space_member(p, str(convo["space_id"]), user["id"])
        if not role:
            return _json_error("not a member of this space", 403)

        search_depth = data.get("search_depth", "basic")
        max_results = int(data.get("max_results", 5))

        t0 = time.monotonic()
        async with httpx.AsyncClient() as client:
            resp = await client.post(
                TAVILY_API_URL,
                json={
                    "api_key": TAVILY_API_KEY,
                    "query": query,
                    "search_depth": search_depth,
                    "max_results": max_results,
                    "include_answer": True,
                    "include_raw_content": True,
                },
                timeout=20,
            )
            resp.raise_for_status()
            payload = resp.json()
        duration_ms = int((time.monotonic() - t0) * 1000)

        row = await p.fetchrow(
            """
            INSERT INTO chat.web_searches
                (conversation_id, message_id, query, search_depth, result_count, response_time_ms, raw_response)
            VALUES ($1::uuid, $2::uuid, $3, $4, $5, $6, $7)
            RETURNING id
            """,
            conversation_id,
            data.get("message_id"),
            query,
            search_depth,
            len(payload.get("results", [])),
            duration_ms,
            _jsonb(payload),
        )

        results = []
        for i, r in enumerate(payload.get("results", []), start=1):
            results.append(
                (
                    row["id"],
                    i,
                    r.get("title"),
                    r.get("url"),
                    r.get("content") or r.get("snippet"),
                    r.get("raw_content"),
                    _jsonb(r),
                )
            )
        if results:
            await p.executemany(
                """
                INSERT INTO chat.web_search_results
                    (search_id, rank, title, url, snippet, content, metadata)
                VALUES ($1, $2, $3, $4, $5, $6, $7)
                """,
                results,
            )

        return _json({"search_id": str(row["id"]), "results": payload.get("results", []), "answer": payload.get("answer")})
    except PermissionError as e:
        return _json_error(str(e), 401)


async def create_memory(request: Request) -> JSONResponse:
    try:
        user = await _get_current_user(request)
        data = await _read_json(request)
        scope = (data.get("scope") or "").strip()
        key = (data.get("key") or "").strip()
        value = data.get("value")
        if scope not in ("global", "user", "space", "conversation"):
            return _json_error("invalid scope")
        if not key or value is None:
            return _json_error("key and value are required")

        p = await get_pool()
        row = await p.fetchrow(
            """
            INSERT INTO chat.memories
                (scope, user_id, space_id, conversation_id, key, value, tags, source_message_id)
            VALUES ($1, $2, $3::uuid, $4::uuid, $5, $6, $7, $8::uuid)
            RETURNING *
            """,
            scope,
            user["id"] if scope in ("user", "space", "conversation") else None,
            data.get("space_id"),
            data.get("conversation_id"),
            key,
            _jsonb(value),
            data.get("tags") or [],
            data.get("source_message_id"),
        )
        return _json(_row(row), status=201)
    except PermissionError as e:
        return _json_error(str(e), 401)


async def list_memories(request: Request) -> JSONResponse:
    try:
        user = await _get_current_user(request)
        p = await get_pool()
        params = request.query_params
        scope = params.get("scope")
        key = params.get("key")
        space_id = params.get("space_id")
        conversation_id = params.get("conversation_id")
        limit = min(int(params.get("limit", "50")), 200)

        where = ["is_active = TRUE"]
        vals = []
        idx = 1

        if scope:
            where.append(f"scope = ${idx}")
            vals.append(scope)
            idx += 1
        if key:
            where.append(f"key = ${idx}")
            vals.append(key)
            idx += 1
        if space_id:
            where.append(f"space_id = ${idx}::uuid")
            vals.append(space_id)
            idx += 1
        if conversation_id:
            where.append(f"conversation_id = ${idx}::uuid")
            vals.append(conversation_id)
            idx += 1

        if scope in ("user", "space", "conversation") or not scope:
            where.append(f"(user_id IS NULL OR user_id = ${idx})")
            vals.append(user["id"])
            idx += 1

        rows = await p.fetch(
            f"SELECT * FROM chat.memories WHERE {' AND '.join(where)} ORDER BY updated_at DESC LIMIT ${idx}",
            *vals,
            limit,
        )
        return _json([_row(r) for r in rows])
    except PermissionError as e:
        return _json_error(str(e), 401)


async def create_completion(request: Request) -> JSONResponse:
    """Generate AI completion for a conversation."""
    request_id = None
    try:
        user = await _get_current_user(request)
        convo_id = request.path_params["conversation_id"]
        data = await _read_json(request)
        print(f"[COMPLETION] convo={convo_id} quality={data.get('quality')} model={data.get('model')} stream={data.get('stream')} keys={list(data.keys())}")
        
        if not OPENAI_API_KEY:
            return _json_error("OPENAI_API_KEY is not configured", 500)
        
        p = await get_pool()
        convo = await p.fetchrow("SELECT * FROM chat.conversations WHERE id = $1::uuid", convo_id)
        if not convo:
            return _json_error("conversation not found", 404)
        
        role = await _require_space_member(p, str(convo["space_id"]), user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        
        # Determine quality settings (matching dashboard pattern)
        quality = (data.get("quality") or "").lower()
        
        if quality in QUALITY_TIERS:
            tier = QUALITY_TIERS[quality]
            model = tier["model"]
            max_completion_tokens = data.get("max_completion_tokens") if "max_completion_tokens" in data else tier["max_completion_tokens"]
            reasoning_effort = tier.get("reasoning_effort")
        else:
            # Always use DEFAULT_MODEL (gpt-5.4) — ignore legacy model stored on conversation
            model = DEFAULT_MODEL
            max_completion_tokens = data.get("max_completion_tokens", 16384)
            reasoning_effort = data.get("reasoning_effort")
        
        # Get conversation history
        messages_rows = await p.fetch(
            """
            SELECT * FROM chat.messages
            WHERE conversation_id = $1::uuid
            ORDER BY ordinal ASC
            """,
            convo_id,
        )
        
        # Create completion request for tracking/cancellation
        request_row = await p.fetchrow(
            """
            INSERT INTO chat.completion_requests
                (conversation_id, user_id, status, quality_tier, model, message_count)
            VALUES ($1::uuid, $2, 'running', $3, $4, $5)
            RETURNING id
            """,
            convo_id,
            user["id"],
            quality or "auto",
            model,
            len(messages_rows),
        )
        request_id = request_row["id"]
        print(f"[COMPLETION] Created request {request_id}")
        
        # Build OpenAI messages array — always inject Athena AI system prompt
        openai_messages = [{"role": "system", "content": ATHENA_SYSTEM_PROMPT}]
        if convo.get("system_prompt"):
            openai_messages.append({"role": "system", "content": convo["system_prompt"]})
        
        for msg in messages_rows:
            if msg["role"] in ("user", "assistant", "system"):
                content = msg.get("content_text") or ""
                if not content:
                    continue

                # Check if this message has image attachments → use vision
                msg_assets = await p.fetch(
                    """
                    SELECT a.* FROM chat.assets a
                    JOIN chat.message_assets ma ON ma.asset_id = a.id
                    WHERE ma.message_id = $1::uuid
                    """,
                    msg["id"],
                )

                image_assets = [
                    a for a in msg_assets
                    if _is_image_content_type(a["content_type"], a["file_name"])
                ]

                if image_assets and msg["role"] == "user":
                    # Build multimodal content array for GPT vision
                    parts = [{"type": "text", "text": content}]
                    for asset in image_assets:
                        fpath = Path(asset["storage_path"])
                        if fpath.exists():
                            img_bytes = fpath.read_bytes()
                            img_b64 = base64.b64encode(img_bytes).decode()
                            ct = asset["content_type"] or "image/png"
                            parts.append({
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{ct};base64,{img_b64}",
                                    "detail": "auto",
                                },
                            })
                    openai_messages.append({"role": "user", "content": parts})
                else:
                    openai_messages.append({"role": msg["role"], "content": content})
        
        # Add user message if provided
        user_message = data.get("message")
        if user_message:
            # Weather detection: inject location context if asking about weather
            if _is_weather_question(user_message):
                location = None
                # Check for GPS coordinates from the app
                gps_lat = data.get("latitude") or data.get("lat")
                gps_lon = data.get("longitude") or data.get("lon") or data.get("lng")
                if gps_lat and gps_lon:
                    location_note = f"[LOCATION] User's GPS coordinates: {gps_lat}, {gps_lon} (source: GPS). Use these coordinates to provide accurate local weather."
                    openai_messages.append({"role": "system", "content": location_note})
                else:
                    # Fall back to IP geolocation
                    client_ip = _get_client_ip(request)
                    geo = await _get_location_from_ip(client_ip)
                    if geo:
                        location_note = (
                            f"[LOCATION] Could not get GPS coordinates. Based on the user's IP address, "
                            f"they appear to be in {geo['city']}, {geo['region']}, {geo['country']} "
                            f"(lat: {geo['lat']}, lon: {geo['lon']}). Source: IP geolocation. "
                            f"IMPORTANT: Tell the user you couldn't get their GPS coordinates but based on their IP address, "
                            f"it looks like they are in {geo['city']}, {geo['region']}."
                        )
                        openai_messages.append({"role": "system", "content": location_note})
                    else:
                        openai_messages.append({"role": "system", "content": "[LOCATION] Could not determine user's location. Ask them to share their location or city name."})
            
            openai_messages.append({"role": "user", "content": user_message})
            # Store user message
            await p.fetchrow(
                """
                INSERT INTO chat.messages
                    (conversation_id, role, content_text, created_by)
                VALUES ($1::uuid, 'user', $2, $3)
                RETURNING *
                """,
                convo_id,
                user_message,
                user["id"],
            )
        
        if not openai_messages:
            await p.execute(
                "UPDATE chat.completion_requests SET status = 'failed', completed_at = now(), error_message = $2 WHERE id = $1::uuid",
                request_id,
                "no messages to complete",
            )
            return _json_error("no messages to complete")
        
        stream = data.get("stream", False)
        
        if stream:
            # Streaming response
            async def generate():
                try:
                    nonlocal request_id
                    async with httpx.AsyncClient() as client:
                        async with client.stream(
                            "POST",
                            OPENAI_API_URL,
                            headers={
                                "Authorization": f"Bearer {OPENAI_API_KEY}",
                                "Content-Type": "application/json",
                            },
                            json={
                                "model": model,
                                "messages": openai_messages,
                                "max_completion_tokens": max_completion_tokens,
                                "stream": True,
                                **(({"reasoning_effort": reasoning_effort} if reasoning_effort else {})),
                            },
                            timeout=120,
                        ) as resp:
                            resp.raise_for_status()
                            full_content = []
                            async for line in resp.aiter_lines():
                                # Check for cancellation
                                cancel_check = await p.fetchrow(
                                    "SELECT status FROM chat.completion_requests WHERE id = $1::uuid",
                                    request_id,
                                )
                                if cancel_check and cancel_check["status"] == "cancelled":
                                    print(f"[COMPLETION] Request {request_id} was cancelled, stopping stream")
                                    yield f"data: {json.dumps({'cancelled': True})}\n\n"
                                    return
                                
                                if line.startswith("data: "):
                                    chunk = line[6:]
                                    if chunk == "[DONE]":
                                        break
                                    try:
                                        data_obj = json.loads(chunk)
                                        delta = data_obj.get("choices", [{}])[0].get("delta", {})
                                        content = delta.get("content", "")
                                        if content:
                                            full_content.append(content)
                                            yield f"data: {json.dumps({'content': content})}\n\n"
                                    except Exception:
                                        pass
                            
                            # Store assistant response
                            if full_content:
                                assistant_text = "".join(full_content)
                                await p.execute(
                                    """
                                    INSERT INTO chat.messages
                                        (conversation_id, role, content_text, created_by, quality_tier)
                                    VALUES ($1::uuid, 'assistant', $2, NULL, $3)
                                    """,
                                    convo_id,
                                    assistant_text,
                                    quality or "auto",
                                )
                                # Auto-generate title after first exchange
                                if convo["title"] in (None, "", "New Conversation", "New chat") and user_message:
                                    title = await _generate_title(user_message, assistant_text)
                                    await p.execute(
                                        "UPDATE chat.conversations SET title = $2, updated_at = now() WHERE id = $1::uuid",
                                        convo_id,
                                        title,
                                    )
                                    yield f"data: {json.dumps({'conversation_title': title})}\n\n"
                                
                                # Mark request as completed
                                await p.execute(
                                    "UPDATE chat.completion_requests SET status = 'completed', completed_at = now() WHERE id = $1::uuid",
                                    request_id,
                                )
                                
                                yield f"data: {json.dumps({'done': True, 'request_id': str(request_id)})}\n\n"
                except Exception as e:
                    import traceback
                    print(f"[COMPLETION STREAM ERROR] {type(e).__name__}: {e}")
                    traceback.print_exc()
                    if request_id:
                        await p.execute(
                            "UPDATE chat.completion_requests SET status = 'failed', completed_at = now(), error_message = $2 WHERE id = $1::uuid",
                            request_id,
                            str(e),
                        )
                    yield f"data: {json.dumps({'error': str(e)})}\n\n"
            
            return StreamingResponse(generate(), media_type="text/event-stream")
        
        else:
            # Non-streaming response
            try:
                async with httpx.AsyncClient() as client:
                    resp = await client.post(
                        OPENAI_API_URL,
                        headers={
                            "Authorization": f"Bearer {OPENAI_API_KEY}",
                            "Content-Type": "application/json",
                        },
                        json={
                            "model": model,
                            "messages": openai_messages,
                            "max_completion_tokens": max_completion_tokens,
                            **(({"reasoning_effort": reasoning_effort} if reasoning_effort else {})),
                        },
                        timeout=120,
                    )
                    resp.raise_for_status()
                    result = resp.json()
                
                assistant_message = result["choices"][0]["message"]["content"]
                usage = result.get("usage", {})
                
                # Store assistant response
                row = await p.fetchrow(
                    """
                    INSERT INTO chat.messages
                        (conversation_id, role, content_text, token_count, created_by, quality_tier)
                    VALUES ($1::uuid, 'assistant', $2, $3, NULL, $4)
                    RETURNING *
                    """,
                    convo_id,
                    assistant_message,
                    usage.get("total_tokens"),
                    quality or "auto",
                )
                
                # Auto-generate title after first exchange
                conversation_title = convo["title"]
                if conversation_title in (None, "", "New Conversation", "New chat") and user_message:
                    conversation_title = await _generate_title(user_message, assistant_message)
                    await p.execute(
                        "UPDATE chat.conversations SET title = $2, updated_at = now() WHERE id = $1::uuid",
                        convo_id,
                        conversation_title,
                    )
                
                # Mark request as completed
                await p.execute(
                    "UPDATE chat.completion_requests SET status = 'completed', completed_at = now() WHERE id = $1::uuid",
                    request_id,
                )
                
                return _json({
                    "message": _row(row),
                    "usage": usage,
                    "model": model,
                    "conversation_title": conversation_title,
                    "request_id": str(request_id),
                })
            except httpx.HTTPStatusError as e:
                import traceback
                print(f"[COMPLETION 502] model={model} quality={quality} reasoning_effort={reasoning_effort}")
                print(f"[COMPLETION 502] OpenAI response: {e.response.text}")
                traceback.print_exc()
                if request_id:
                    await p.execute(
                        "UPDATE chat.completion_requests SET status = 'failed', completed_at = now(), error_message = $2 WHERE id = $1::uuid",
                        request_id,
                        e.response.text,
                    )
                return _json_error(f"OpenAI API error: {e.response.text}", 502)
    
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback
        print(f"[COMPLETION ERROR] {type(e).__name__}: {str(e)}")
        traceback.print_exc()
        if request_id:
            p = await get_pool()
            await p.execute(
                "UPDATE chat.completion_requests SET status = 'failed', completed_at = now(), error_message = $2 WHERE id = $1::uuid",
                request_id,
                str(e),
            )
        return _json_error(f"Completion error: {str(e)}", 500)


# ══════════════════════════════════════════════════════════════
# SPACE KNOWLEDGE BASE - RAG & Semantic Search
# ══════════════════════════════════════════════════════════════

async def _generate_embedding(text: str) -> list[float]:
    """Generate embedding for text using OpenAI's text-embedding-3-large."""
    if not OPENAI_API_KEY:
        raise RuntimeError("OPENAI_API_KEY not configured")
    
    async with httpx.AsyncClient() as client:
        resp = await client.post(
            "https://api.openai.com/v1/embeddings",
            headers={
                "Authorization": f"Bearer {OPENAI_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": "text-embedding-3-large",
                "input": text[:8000],  # Limit to avoid token overflow
            },
            timeout=30,
        )
        resp.raise_for_status()
        return resp.json()["data"][0]["embedding"]


def _chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[dict]:
    """Chunk text into overlapping segments for better embedding coverage."""
    if not text or len(text) < chunk_size:
        return [{"text": text, "start": 0, "end": len(text), "index": 0}]
    
    chunks = []
    start = 0
    index = 0
    
    while start < len(text):
        end = start + chunk_size
        # Try to break at sentence boundary
        if end < len(text):
            period_pos = text.rfind(". ", start, end)
            if period_pos > start + chunk_size // 2:
                end = period_pos + 1
        
        chunk_text = text[start:end].strip()
        if chunk_text:
            chunks.append({
                "text": chunk_text,
                "start": start,
                "end": end,
                "index": index,
            })
            index += 1
        
        start = end - overlap
    
    return chunks


async def add_space_document(request: Request) -> JSONResponse:
    """Add a document (email, file, note, etc.) to a space's knowledge base.
    
    Body JSON:
      source_type — email, file, note, url, attachment, case, medical_record, other
      title — document title
      content_text — plain text content
      content_html — optional HTML content
      summary — optional summary
      source_system — wdws, m365, external, etc.
      source_id — external ID
      source_uri — external URI
      file_name — file name if applicable
      content_type — MIME type
      author_name, author_email — author info
      document_date — document timestamp
      tags — array of tags
      metadata — additional metadata
      auto_chunk — if true, automatically chunk and embed (default: true)
    """
    try:
        user = await _get_current_user(request)
        space_id = request.path_params["space_id"]
        data = await _read_json(request)
        
        source_type = data.get("source_type", "other")
        title = (data.get("title") or "").strip()
        if not title:
            return _json_error("title is required")
        
        content_text = data.get("content_text", "")
        if not content_text:
            return _json_error("content_text is required")
        
        p = await get_pool()
        
        # Verify user has access to space
        role = await _require_space_member(p, space_id, user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        
        # Create document
        doc_row = await p.fetchrow(
            """
            INSERT INTO chat.space_documents
                (space_id, source_type, title, content_text, content_html, summary,
                 source_system, source_id, source_uri, file_name, content_type,
                 author_name, author_email, document_date, tags, metadata,
                 status, created_by)
            VALUES ($1::uuid, $2, $3, $4, $5, $6, $7, $8, $9, $10, $11, $12, $13, $14, $15, $16, $17, $18)
            RETURNING *
            """,
            space_id,
            source_type,
            title,
            content_text,
            data.get("content_html"),
            data.get("summary"),
            data.get("source_system"),
            data.get("source_id"),
            data.get("source_uri"),
            data.get("file_name"),
            data.get("content_type"),
            data.get("author_name"),
            data.get("author_email"),
            data.get("document_date"),
            data.get("tags", []),
            _jsonb(data.get("metadata") or {}),
            "pending",
            user["id"],
        )
        
        doc_id = doc_row["id"]
        auto_chunk = data.get("auto_chunk", True)
        
        # Auto-chunk and embed if requested
        if auto_chunk and OPENAI_API_KEY:
            try:
                await p.execute(
                    "UPDATE chat.space_documents SET status = 'processing' WHERE id = $1::uuid",
                    doc_id,
                )
                
                # Chunk the document
                chunks = _chunk_text(content_text, chunk_size=1000, overlap=200)
                print(f"[SPACE KB] Document {doc_id} chunked into {len(chunks)} pieces")
                
                # Store chunks and generate embeddings
                for chunk_data in chunks:
                    chunk_row = await p.fetchrow(
                        """
                        INSERT INTO chat.space_document_chunks
                            (document_id, space_id, chunk_index, chunk_text, start_offset, end_offset, token_count)
                        VALUES ($1::uuid, $2::uuid, $3, $4, $5, $6, $7)
                        RETURNING id
                        """,
                        doc_id,
                        space_id,
                        chunk_data["index"],
                        chunk_data["text"],
                        chunk_data["start"],
                        chunk_data["end"],
                        len(chunk_data["text"].split()),  # Rough token count
                    )
                    
                    # Generate embedding
                    embedding = await _generate_embedding(chunk_data["text"])
                    
                    # Store embedding
                    await p.execute(
                        """
                        INSERT INTO chat.space_document_embeddings
                            (chunk_id, space_id, embedding, embedding_model_id)
                        VALUES ($1::uuid, $2::uuid, $3, 1)
                        """,
                        chunk_row["id"],
                        space_id,
                        embedding,
                    )
                
                # Mark as indexed
                await p.execute(
                    """
                    UPDATE chat.space_documents
                    SET status = 'indexed', last_indexed_at = now()
                    WHERE id = $1::uuid
                    """,
                    doc_id,
                )
                
                print(f"[SPACE KB] Document {doc_id} indexed successfully")
            
            except Exception as e:
                import traceback
                print(f"[SPACE KB ERROR] Failed to index document {doc_id}: {e}")
                traceback.print_exc()
                await p.execute(
                    """
                    UPDATE chat.space_documents
                    SET status = 'failed', processing_error = $2
                    WHERE id = $1::uuid
                    """,
                    doc_id,
                    str(e),
                )
        
        return _json(_row(doc_row), status=201)
    
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback
        print(f"[SPACE KB ERROR] {type(e).__name__}: {e}")
        traceback.print_exc()
        return _json_error(f"Add document error: {str(e)}", 500)


async def search_space_knowledge(request: Request) -> JSONResponse:
    """Semantic search within a space's knowledge base.
    
    Query params:
      q or query — search query
      search_type — semantic, fulltext, hybrid (default: hybrid)
      limit — max results (default: 10, max: 50)
      min_score — minimum similarity score 0-1 (default: 0.7)
      source_type — filter by source type
      tags — filter by tags (comma-separated)
    """
    try:
        user = await _get_current_user(request)
        space_id = request.path_params["space_id"]
        
        query = request.query_params.get("q") or request.query_params.get("query", "").strip()
        if not query:
            return _json_error("query parameter required")
        
        search_type = request.query_params.get("search_type", "hybrid")
        if search_type not in ("semantic", "fulltext", "hybrid"):
            return _json_error("search_type must be semantic, fulltext, or hybrid")
        
        limit = min(int(request.query_params.get("limit", 10)), 50)
        min_score = float(request.query_params.get("min_score", 0.7))
        source_type_filter = request.query_params.get("source_type")
        tags_filter = request.query_params.get("tags", "").split(",") if request.query_params.get("tags") else None
        
        p = await get_pool()
        
        # Verify user has access to space
        role = await _require_space_member(p, space_id, user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        
        start_time = time.time()
        results = []
        
        if search_type in ("semantic", "hybrid"):
            # Generate query embedding
            query_embedding = await _generate_embedding(query)
            
            # Semantic search with space isolation
            semantic_sql = """
                SELECT
                    c.id AS chunk_id,
                    c.chunk_text,
                    c.chunk_index,
                    d.id AS document_id,
                    d.title,
                    d.source_type,
                    d.author_name,
                    d.document_date,
                    d.tags,
                    1 - (e.embedding <=> $1::halfvec) AS similarity_score
                FROM chat.space_document_embeddings e
                JOIN chat.space_document_chunks c ON c.id = e.chunk_id
                JOIN chat.space_documents d ON d.id = c.document_id
                WHERE e.space_id = $2::uuid
                  AND d.status = 'indexed'
            """
            params = [query_embedding, space_id]
            param_idx = 3
            
            if source_type_filter:
                semantic_sql += f" AND d.source_type = ${param_idx}"
                params.append(source_type_filter)
                param_idx += 1
            
            if tags_filter:
                semantic_sql += f" AND d.tags && ${param_idx}"
                params.append(tags_filter)
                param_idx += 1
            
            semantic_sql += f"""
                ORDER BY e.embedding <=> $1::halfvec
                LIMIT ${param_idx}
            """
            params.append(limit)
            
            rows = await p.fetch(semantic_sql, *params)
            
            for row in rows:
                score = row["similarity_score"]
                if score >= min_score:
                    results.append({
                        "document_id": str(row["document_id"]),
                        "title": row["title"],
                        "source_type": row["source_type"],
                        "chunk_text": row["chunk_text"][:500],  # Preview
                        "chunk_index": row["chunk_index"],
                        "author_name": row["author_name"],
                        "document_date": row["document_date"],
                        "tags": row["tags"],
                        "score": float(score),
                        "search_type": "semantic",
                    })
        
        elif search_type == "fulltext":
            # Full-text search
            fulltext_sql = """
                SELECT
                    d.id AS document_id,
                    d.title,
                    d.source_type,
                    d.content_text,
                    d.author_name,
                    d.document_date,
                    d.tags,
                    ts_rank(
                        to_tsvector('english',
                            coalesce(d.title, '') || ' ' ||
                            coalesce(d.content_text, '') || ' ' ||
                            coalesce(d.author_name, '') || ' ' ||
                            coalesce(array_to_string(d.tags, ' '), '')
                        ),
                        plainto_tsquery('english', $1)
                    ) AS rank
                FROM chat.space_documents d
                WHERE d.space_id = $2::uuid
                  AND d.status = 'indexed'
                  AND to_tsvector('english',
                        coalesce(d.title, '') || ' ' ||
                        coalesce(d.content_text, '') || ' ' ||
                        coalesce(d.author_name, '') || ' ' ||
                        coalesce(array_to_string(d.tags, ' '), '')
                    ) @@ plainto_tsquery('english', $1)
            """
            params = [query, space_id]
            param_idx = 3
            
            if source_type_filter:
                fulltext_sql += f" AND d.source_type = ${param_idx}"
                params.append(source_type_filter)
                param_idx += 1
            
            if tags_filter:
                fulltext_sql += f" AND d.tags && ${param_idx}"
                params.append(tags_filter)
                param_idx += 1
            
            fulltext_sql += f"""
                ORDER BY rank DESC
                LIMIT ${param_idx}
            """
            params.append(limit)
            
            rows = await p.fetch(fulltext_sql, *params)
            
            for row in rows:
                results.append({
                    "document_id": str(row["document_id"]),
                    "title": row["title"],
                    "source_type": row["source_type"],
                    "content_preview": row["content_text"][:500] if row["content_text"] else "",
                    "author_name": row["author_name"],
                    "document_date": row["document_date"],
                    "tags": row["tags"],
                    "score": float(row["rank"]),
                    "search_type": "fulltext",
                })
        
        execution_time_ms = int((time.time() - start_time) * 1000)
        
        # Log search
        await p.execute(
            """
            INSERT INTO chat.space_searches
                (space_id, user_id, query, search_type, result_count, execution_time_ms)
            VALUES ($1::uuid, $2, $3, $4, $5, $6)
            """,
            space_id,
            user["id"],
            query,
            search_type,
            len(results),
            execution_time_ms,
        )
        
        print(f"[SPACE SEARCH] space={space_id} query={query!r} type={search_type} results={len(results)} time={execution_time_ms}ms")
        
        return _json({
            "query": query,
            "search_type": search_type,
            "results": results,
            "count": len(results),
            "execution_time_ms": execution_time_ms,
        })
    
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback
        print(f"[SPACE SEARCH ERROR] {type(e).__name__}: {e}")
        traceback.print_exc()
        return _json_error(f"Search error: {str(e)}", 500)


async def list_space_documents(request: Request) -> JSONResponse:
    """List documents in a space's knowledge base."""
    try:
        user = await _get_current_user(request)
        space_id = request.path_params["space_id"]
        
        limit = min(int(request.query_params.get("limit", 50)), 200)
        offset = int(request.query_params.get("offset", 0))
        source_type = request.query_params.get("source_type")
        status = request.query_params.get("status", "indexed")
        
        p = await get_pool()
        
        # Verify user has access
        role = await _require_space_member(p, space_id, user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        
        # Build query
        where_clauses = ["space_id = $1::uuid"]
        params = [space_id]
        param_idx = 2
        
        if source_type:
            where_clauses.append(f"source_type = ${param_idx}")
            params.append(source_type)
            param_idx += 1
        
        if status:
            where_clauses.append(f"status = ${param_idx}")
            params.append(status)
            param_idx += 1
        
        sql = f"""
            SELECT * FROM chat.space_documents
            WHERE {' AND '.join(where_clauses)}
            ORDER BY created_at DESC
            LIMIT ${param_idx} OFFSET ${param_idx + 1}
        """
        params.extend([limit, offset])
        
        rows = await p.fetch(sql, *params)
        
        return _json([_row(r) for r in rows])
    
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback
        print(f"[LIST DOCS ERROR] {type(e).__name__}: {e}")
        traceback.print_exc()
        return _json_error(f"List documents error: {str(e)}", 500)


async def get_space_document(request: Request) -> JSONResponse:
    """Get a specific document from a space's knowledge base."""
    try:
        user = await _get_current_user(request)
        space_id = request.path_params["space_id"]
        document_id = request.path_params["document_id"]
        
        include_chunks = request.query_params.get("include_chunks", "false").lower() == "true"
        
        p = await get_pool()
        
        # Verify user has access
        role = await _require_space_member(p, space_id, user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        
        # Get document
        doc = await p.fetchrow(
            "SELECT * FROM chat.space_documents WHERE id = $1::uuid AND space_id = $2::uuid",
            document_id,
            space_id,
        )
        
        if not doc:
            return _json_error("document not found", 404)
        
        result = _row(doc)
        
        # Include chunks if requested
        if include_chunks:
            chunks = await p.fetch(
                """
                SELECT * FROM chat.space_document_chunks
                WHERE document_id = $1::uuid
                ORDER BY chunk_index
                """,
                document_id,
            )
            result["chunks"] = [_row(c) for c in chunks]
        
        return _json(result)
    
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback
        print(f"[GET DOC ERROR] {type(e).__name__}: {e}")
        traceback.print_exc()
        return _json_error(f"Get document error: {str(e)}", 500)


async def create_space_note(request: Request) -> JSONResponse:
    """Create a quick note in a space.
    
    Body JSON:
      title — note title (optional)
      content — note content
      format — markdown, html, plain (default: markdown)
      tags — array of tags
      color — color code
      is_pinned — pin to top
    """
    try:
        user = await _get_current_user(request)
        space_id = request.path_params["space_id"]
        data = await _read_json(request)
        
        content = (data.get("content") or "").strip()
        if not content:
            return _json_error("content is required")
        
        format_type = data.get("format", "markdown")
        if format_type not in ("markdown", "html", "plain"):
            return _json_error("format must be markdown, html, or plain")
        
        p = await get_pool()
        
        # Verify user has access
        role = await _require_space_member(p, space_id, user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        
        row = await p.fetchrow(
            """
            INSERT INTO chat.space_notes
                (space_id, title, content, format, tags, color, is_pinned, created_by)
            VALUES ($1::uuid, $2, $3, $4, $5, $6, $7, $8)
            RETURNING *
            """,
            space_id,
            data.get("title"),
            content,
            format_type,
            data.get("tags", []),
            data.get("color"),
            bool(data.get("is_pinned", False)),
            user["id"],
        )
        
        return _json(_row(row), status=201)
    
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback
        print(f"[CREATE NOTE ERROR] {type(e).__name__}: {e}")
        traceback.print_exc()
        return _json_error(f"Create note error: {str(e)}", 500)


async def list_space_notes(request: Request) -> JSONResponse:
    """List notes in a space."""
    try:
        user = await _get_current_user(request)
        space_id = request.path_params["space_id"]
        
        p = await get_pool()
        
        # Verify user has access
        role = await _require_space_member(p, space_id, user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        
        rows = await p.fetch(
            """
            SELECT * FROM chat.space_notes
            WHERE space_id = $1::uuid
            ORDER BY is_pinned DESC, created_at DESC
            """,
            space_id,
        )
        
        return _json([_row(r) for r in rows])
    
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback
        print(f"[LIST NOTES ERROR] {type(e).__name__}: {e}")
        traceback.print_exc()
        return _json_error(f"List notes error: {str(e)}", 500)


async def get_space_stats(request: Request) -> JSONResponse:
    """Get knowledge base statistics for a space."""
    try:
        user = await _get_current_user(request)
        space_id = request.path_params["space_id"]
        
        p = await get_pool()
        
        # Verify user has access
        role = await _require_space_member(p, space_id, user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        
        # Get stats from view
        stats = await p.fetchrow(
            "SELECT * FROM chat.v_space_document_stats WHERE space_id = $1::uuid",
            space_id,
        )
        
        # Get note count
        note_count = await p.fetchval(
            "SELECT COUNT(*) FROM chat.space_notes WHERE space_id = $1::uuid",
            space_id,
        )
        
        # Get recent searches
        recent_searches = await p.fetch(
            """
            SELECT query, search_type, result_count, created_at
            FROM chat.space_searches
            WHERE space_id = $1::uuid
            ORDER BY created_at DESC
            LIMIT 10
            """,
            space_id,
        )
        
        return _json({
            "space_id": space_id,
            "documents": _row(stats) if stats else {},
            "notes_count": note_count,
            "recent_searches": [_row(s) for s in recent_searches],
        })
    
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback
        print(f"[SPACE STATS ERROR] {type(e).__name__}: {e}")
        traceback.print_exc()
        return _json_error(f"Space stats error: {str(e)}", 500)


# ══════════════════════════════════════════════════════════════
# ONEDRIVE INTEGRATION
# ══════════════════════════════════════════════════════════════

async def list_onedrive_targets(request: Request) -> JSONResponse:
    """List configured OneDrive sync targets."""
    try:
        user = await _get_current_user(request)
        p = await get_pool()
        rows = await p.fetch("""
            SELECT id, drive_type, drive_owner, folder_path, is_active, sync_mode,
                   label, total_files, total_chunks,
                   pg_size_pretty(total_bytes) as total_size,
                   last_sync_at, last_sync_error,
                   file_extensions, created_at
            FROM chat.onedrive_sync_targets ORDER BY id
        """)
        return _json([_row(r) for r in rows])
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        return _json_error(f"Error: {e}", 500)


async def list_onedrive_files(request: Request) -> JSONResponse:
    """List indexed OneDrive files.
    
    Query params:
      target_id — filter by sync target
      status — filter by status (indexed, pending, failed, etc.)
      extension — filter by file extension
      q — search file names
      limit — max results (default: 50, max: 200)
      offset — pagination offset
    """
    try:
        user = await _get_current_user(request)
        p = await get_pool()
        
        target_id = request.query_params.get("target_id")
        status_filter = request.query_params.get("status")
        ext_filter = request.query_params.get("extension")
        q = request.query_params.get("q", "").strip()
        limit = min(int(request.query_params.get("limit", 50)), 200)
        offset = int(request.query_params.get("offset", 0))
        
        sql = """
            SELECT f.id, f.file_name, f.file_path, f.file_extension,
                   f.mime_type, f.size_bytes, f.status,
                   f.created_by, f.last_modified_by, f.web_url,
                   f.graph_modified_at, f.indexed_at,
                   f.content_summary, f.extraction_method,
                   t.label as target_label
            FROM chat.onedrive_files f
            LEFT JOIN chat.onedrive_sync_targets t ON t.id = f.sync_target_id
            WHERE f.status != 'deleted'
        """
        params = []
        idx = 1
        
        if target_id:
            sql += f" AND f.sync_target_id = ${idx}::bigint"
            params.append(int(target_id))
            idx += 1
        if status_filter:
            sql += f" AND f.status = ${idx}"
            params.append(status_filter)
            idx += 1
        if ext_filter:
            sql += f" AND f.file_extension = ${idx}"
            params.append(ext_filter)
            idx += 1
        if q:
            sql += f" AND f.file_name ILIKE ${idx}"
            params.append(f"%{q}%")
            idx += 1
        
        sql += f" ORDER BY f.graph_modified_at DESC NULLS LAST LIMIT ${idx} OFFSET ${idx+1}"
        params.extend([limit, offset])
        
        rows = await p.fetch(sql, *params)
        
        # Get total count
        count_sql = "SELECT COUNT(*) FROM chat.onedrive_files WHERE status != 'deleted'"
        total = await p.fetchval(count_sql)
        
        return _json({
            "files": [_row(r) for r in rows],
            "total": total,
            "limit": limit,
            "offset": offset,
        })
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        return _json_error(f"Error: {e}", 500)


async def search_onedrive(request: Request) -> JSONResponse:
    """Semantic/hybrid search across all indexed OneDrive content.
    
    Query params:
      q — search query (required)
      search_type — semantic, fulltext, hybrid (default: hybrid)
      limit — max results (default: 10, max: 50)
      min_score — minimum similarity (default: 0.65)
      extension — filter by file extension
    """
    try:
        user = await _get_current_user(request)
        p = await get_pool()
        
        query = (request.query_params.get("q") or request.query_params.get("query", "")).strip()
        if not query:
            return _json_error("q parameter required")
        
        search_type = request.query_params.get("search_type", "hybrid")
        limit = min(int(request.query_params.get("limit", 10)), 50)
        min_score = float(request.query_params.get("min_score", 0.65))
        ext_filter = request.query_params.get("extension")
        
        start_time = time.time()
        results = []
        
        if search_type in ("semantic", "hybrid"):
            query_embedding = await _generate_embedding(query)
            sql = """
                SELECT
                    c.id AS chunk_id,
                    c.chunk_text,
                    c.chunk_index,
                    c.page_number,
                    c.sheet_name,
                    f.id AS file_id,
                    f.file_name,
                    f.file_path,
                    f.file_extension,
                    f.web_url,
                    f.created_by,
                    f.last_modified_by,
                    f.graph_modified_at,
                    1 - (c.embedding <=> $1::halfvec) AS similarity
                FROM chat.onedrive_file_chunks c
                JOIN chat.onedrive_files f ON f.id = c.file_id
                WHERE f.status = 'indexed'
                  AND c.embedding IS NOT NULL
            """
            params = [query_embedding]
            idx = 2
            
            if ext_filter:
                sql += f" AND f.file_extension = ${idx}"
                params.append(ext_filter)
                idx += 1
            
            sql += f" ORDER BY c.embedding <=> $1::halfvec LIMIT ${idx}"
            params.append(limit * 2)  # fetch extra, filter by min_score
            
            rows = await p.fetch(sql, *params)
            for row in rows:
                score = float(row["similarity"])
                if score >= min_score:
                    results.append({
                        "file_id": str(row["file_id"]),
                        "file_name": row["file_name"],
                        "file_path": row["file_path"],
                        "file_extension": row["file_extension"],
                        "web_url": row["web_url"],
                        "chunk_text": row["chunk_text"][:500],
                        "chunk_index": row["chunk_index"],
                        "page_number": row["page_number"],
                        "sheet_name": row["sheet_name"],
                        "created_by": row["created_by"],
                        "last_modified_by": row["last_modified_by"],
                        "modified_at": row["graph_modified_at"],
                        "score": score,
                        "search_type": "semantic",
                    })
        
        if search_type in ("fulltext", "hybrid"):
            ft_sql = """
                SELECT
                    f.id AS file_id,
                    f.file_name,
                    f.file_path,
                    f.file_extension,
                    f.web_url,
                    f.created_by,
                    f.last_modified_by,
                    f.graph_modified_at,
                    ts_rank(
                        to_tsvector('english', coalesce(f.file_name, '') || ' ' || coalesce(f.content_text, '')),
                        plainto_tsquery('english', $1)
                    ) AS rank
                FROM chat.onedrive_files f
                WHERE f.status = 'indexed'
                  AND to_tsvector('english', coalesce(f.file_name, '') || ' ' || coalesce(f.content_text, ''))
                      @@ plainto_tsquery('english', $1)
            """
            ft_params = [query]
            ft_idx = 2
            
            if ext_filter:
                ft_sql += f" AND f.file_extension = ${ft_idx}"
                ft_params.append(ext_filter)
                ft_idx += 1
            
            ft_sql += f" ORDER BY rank DESC LIMIT ${ft_idx}"
            ft_params.append(limit)
            
            rows = await p.fetch(ft_sql, *ft_params)
            for row in rows:
                # Deduplicate with semantic results
                file_id = str(row["file_id"])
                if not any(r["file_id"] == file_id for r in results):
                    results.append({
                        "file_id": file_id,
                        "file_name": row["file_name"],
                        "file_path": row["file_path"],
                        "file_extension": row["file_extension"],
                        "web_url": row["web_url"],
                        "created_by": row["created_by"],
                        "last_modified_by": row["last_modified_by"],
                        "modified_at": row["graph_modified_at"],
                        "score": float(row["rank"]),
                        "search_type": "fulltext",
                    })
        
        # Sort by score desc, truncate to limit
        results.sort(key=lambda r: r["score"], reverse=True)
        results = results[:limit]
        
        execution_time_ms = int((time.time() - start_time) * 1000)
        print(f"[ONEDRIVE SEARCH] query={query!r} type={search_type} results={len(results)} time={execution_time_ms}ms")
        
        return _json({
            "query": query,
            "search_type": search_type,
            "results": results,
            "count": len(results),
            "execution_time_ms": execution_time_ms,
        })
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return _json_error(f"Search error: {e}", 500)


async def list_space_onedrive_links(request: Request) -> JSONResponse:
    """List OneDrive files/targets linked to a space."""
    try:
        user = await _get_current_user(request)
        space_id = request.path_params["space_id"]
        p = await get_pool()
        
        role = await _require_space_member(p, space_id, user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        
        rows = await p.fetch("""
            SELECT * FROM chat.v_space_onedrive_files
            WHERE space_id = $1::uuid
            ORDER BY file_name
        """, space_id)
        
        return _json([_row(r) for r in rows])
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        return _json_error(f"Error: {e}", 500)


async def link_onedrive_to_space(request: Request) -> JSONResponse:
    """Link a OneDrive file or sync target to a space.
    
    Body JSON:
      file_id — UUID of a specific indexed file
      sync_target_id — ID of a sync target (links all files from that target)
    """
    try:
        user = await _get_current_user(request)
        space_id = request.path_params["space_id"]
        body = await _read_json(request)
        p = await get_pool()
        
        role = await _require_space_member(p, space_id, user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        
        file_id = body.get("file_id")
        sync_target_id = body.get("sync_target_id")
        
        if not file_id and not sync_target_id:
            return _json_error("file_id or sync_target_id required")
        
        if file_id:
            # Verify file exists
            exists = await p.fetchval(
                "SELECT 1 FROM chat.onedrive_files WHERE id = $1::uuid",
                file_id
            )
            if not exists:
                return _json_error("file not found", 404)
            
            link_id = await p.fetchval("""
                INSERT INTO chat.space_onedrive_links (space_id, file_id, linked_by)
                VALUES ($1::uuid, $2::uuid, $3)
                ON CONFLICT (space_id, file_id) DO NOTHING
                RETURNING id
            """, space_id, file_id, user["id"])
        else:
            # Verify target exists
            exists = await p.fetchval(
                "SELECT 1 FROM chat.onedrive_sync_targets WHERE id = $1",
                int(sync_target_id)
            )
            if not exists:
                return _json_error("sync target not found", 404)
            
            link_id = await p.fetchval("""
                INSERT INTO chat.space_onedrive_links (space_id, sync_target_id, linked_by)
                VALUES ($1::uuid, $2, $3)
                ON CONFLICT (space_id, sync_target_id) DO NOTHING
                RETURNING id
            """, space_id, int(sync_target_id), user["id"])
        
        return _json({"linked": link_id is not None, "link_id": link_id})
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        return _json_error(f"Error: {e}", 500)


async def unlink_onedrive_from_space(request: Request) -> JSONResponse:
    """Remove a OneDrive link from a space.
    
    Query params:
      file_id — unlink a specific file
      sync_target_id — unlink a sync target
    """
    try:
        user = await _get_current_user(request)
        space_id = request.path_params["space_id"]
        p = await get_pool()
        
        role = await _require_space_member(p, space_id, user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        
        file_id = request.query_params.get("file_id")
        sync_target_id = request.query_params.get("sync_target_id")
        
        if file_id:
            await p.execute("""
                DELETE FROM chat.space_onedrive_links
                WHERE space_id = $1::uuid AND file_id = $2::uuid
            """, space_id, file_id)
        elif sync_target_id:
            await p.execute("""
                DELETE FROM chat.space_onedrive_links
                WHERE space_id = $1::uuid AND sync_target_id = $2
            """, space_id, int(sync_target_id))
        else:
            return _json_error("file_id or sync_target_id required")
        
        return _json({"unlinked": True})
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        return _json_error(f"Error: {e}", 500)


async def search_space_onedrive(request: Request) -> JSONResponse:
    """Search OneDrive content that is linked to a specific space.
    
    Query params:
      q — search query (required)
      search_type — semantic, fulltext, hybrid (default: hybrid)
      limit — max results (default: 10, max: 50)
      min_score — minimum similarity (default: 0.65)
    """
    try:
        user = await _get_current_user(request)
        space_id = request.path_params["space_id"]
        p = await get_pool()
        
        role = await _require_space_member(p, space_id, user["id"])
        if not role:
            return _json_error("not a member of this space", 403)
        
        query = (request.query_params.get("q") or request.query_params.get("query", "")).strip()
        if not query:
            return _json_error("q parameter required")
        
        search_type = request.query_params.get("search_type", "hybrid")
        limit = min(int(request.query_params.get("limit", 10)), 50)
        min_score = float(request.query_params.get("min_score", 0.65))
        
        start_time = time.time()
        results = []
        
        # Get the set of file IDs linked to this space
        # (direct file links + all files from linked sync targets)
        linked_file_ids_sql = """
            SELECT f.id FROM chat.onedrive_files f
            JOIN chat.space_onedrive_links sol ON 
                (sol.file_id = f.id OR sol.sync_target_id = f.sync_target_id)
            WHERE sol.space_id = $1::uuid
              AND f.status = 'indexed'
        """
        
        if search_type in ("semantic", "hybrid"):
            query_embedding = await _generate_embedding(query)
            sql = f"""
                SELECT
                    c.chunk_text,
                    c.chunk_index,
                    c.page_number,
                    f.id AS file_id,
                    f.file_name,
                    f.file_path,
                    f.web_url,
                    f.file_extension,
                    f.last_modified_by,
                    f.graph_modified_at,
                    1 - (c.embedding <=> $1::halfvec) AS similarity
                FROM chat.onedrive_file_chunks c
                JOIN chat.onedrive_files f ON f.id = c.file_id
                WHERE f.id IN ({linked_file_ids_sql})
                  AND c.embedding IS NOT NULL
                ORDER BY c.embedding <=> $1::halfvec
                LIMIT $2
            """
            rows = await p.fetch(sql, query_embedding, space_id, limit * 2)
            for row in rows:
                score = float(row["similarity"])
                if score >= min_score:
                    results.append({
                        "file_id": str(row["file_id"]),
                        "file_name": row["file_name"],
                        "file_path": row["file_path"],
                        "web_url": row["web_url"],
                        "chunk_text": row["chunk_text"][:500],
                        "chunk_index": row["chunk_index"],
                        "page_number": row["page_number"],
                        "last_modified_by": row["last_modified_by"],
                        "modified_at": row["graph_modified_at"],
                        "score": score,
                        "search_type": "semantic",
                    })
        
        results.sort(key=lambda r: r["score"], reverse=True)
        results = results[:limit]
        
        execution_time_ms = int((time.time() - start_time) * 1000)
        
        return _json({
            "query": query,
            "search_type": search_type,
            "space_id": space_id,
            "results": results,
            "count": len(results),
            "execution_time_ms": execution_time_ms,
        })
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return _json_error(f"Search error: {e}", 500)


async def onedrive_webhook(request: Request) -> JSONResponse:
    """Receive Graph API change notifications for OneDrive.
    
    Handles:
      - Validation requests (GET with validationToken)
      - Change notifications (POST with notification payload)
    """
    # Validation handshake — Graph sends GET with validationToken
    validation_token = request.query_params.get("validationToken")
    if validation_token:
        from starlette.responses import PlainTextResponse
        return PlainTextResponse(validation_token)
    
    try:
        body = await request.json()
        notifications = body.get("value", [])
        
        for notif in notifications:
            # Verify client state
            if notif.get("clientState") != "wdws-onedrive-sync":
                continue
            
            resource = notif.get("resource", "")
            change_type = notif.get("changeType", "")
            subscription_id = notif.get("subscriptionId", "")
            
            print(f"[ONEDRIVE WEBHOOK] change={change_type} resource={resource} sub={subscription_id}")
            
            # Find the sync target for this subscription
            p = await get_pool()
            target_id = await p.fetchval("""
                SELECT sync_target_id FROM chat.onedrive_webhooks
                WHERE subscription_id = $1 AND is_active = true
            """, subscription_id)
            
            if target_id:
                # Trigger an async delta sync for this target
                # In production, this would push to a task queue.
                # For now, we just update the target's last_sync_error to signal
                # the daemon to prioritize it on next cycle.
                await p.execute("""
                    UPDATE chat.onedrive_sync_targets 
                    SET last_sync_error = 'webhook_triggered'
                    WHERE id = $1
                """, target_id)
        
        return _json({"status": "ok"})
    except Exception as e:
        print(f"[ONEDRIVE WEBHOOK ERROR] {e}")
        return _json({"status": "ok"})  # Always return 200 to Graph


async def get_onedrive_stats(request: Request) -> JSONResponse:
    """Get overall OneDrive sync statistics."""
    try:
        user = await _get_current_user(request)
        p = await get_pool()
        
        targets = await p.fetch(
            "SELECT * FROM chat.v_onedrive_sync_status ORDER BY target_id"
        )
        
        totals = await p.fetchrow("""
            SELECT
                COUNT(*) FILTER (WHERE status = 'indexed') AS indexed_files,
                COUNT(*) FILTER (WHERE status = 'pending') AS pending_files,
                COUNT(*) FILTER (WHERE status = 'failed') AS failed_files,
                COUNT(*) FILTER (WHERE status = 'skipped') AS skipped_files,
                pg_size_pretty(SUM(size_bytes) FILTER (WHERE status != 'deleted')) AS total_file_size
            FROM chat.onedrive_files
        """)
        
        chunk_count = await p.fetchval(
            "SELECT COUNT(*) FROM chat.onedrive_file_chunks"
        )

        webhook_count = await p.fetchval(
            "SELECT COUNT(*) FROM chat.onedrive_webhooks WHERE is_active = true"
        )

        ext_breakdown = await p.fetch("""
            SELECT file_extension, COUNT(*) as cnt,
                   pg_size_pretty(SUM(size_bytes)) as total_size
            FROM chat.onedrive_files WHERE status != 'deleted'
            GROUP BY file_extension ORDER BY cnt DESC LIMIT 20
        """)
        
        recent_runs = await p.fetch("""
            SELECT r.*, t.label
            FROM chat.onedrive_sync_runs r
            LEFT JOIN chat.onedrive_sync_targets t ON t.id = r.sync_target_id
            ORDER BY r.started_at DESC LIMIT 10
        """)
        
        return _json({
            "targets": [_row(r) for r in targets],
            "totals": _row(totals) if totals else {},
            "total_chunks": chunk_count,
            "webhook_count": webhook_count,
            "extension_breakdown": [_row(r) for r in ext_breakdown],
            "recent_runs": [_row(r) for r in recent_runs],
        })
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        return _json_error(f"Error: {e}", 500)


async def onedrive_webhooks(request: Request) -> JSONResponse:
    """List or create OneDrive webhook subscriptions."""
    try:
        user = await _get_current_user(request)
        p = await get_pool()

        if request.method == "POST":
            # Trigger webhook registration via the sync module
            import subprocess, sys
            result = subprocess.run(
                [sys.executable, "-u", "-c",
                 "import asyncio,sys;sys.path.insert(0,'/opt/wdws');"
                 "from onedrive_sync import GraphClient,ensure_webhook_subscriptions,CHAT_DATABASE_URL;"
                 "import asyncpg;"
                 "asyncio.run((lambda: asyncio.ensure_future("
                 "  (async_run:=lambda: None)() or asyncio.sleep(0)"
                 "))())" ],
                capture_output=True, text=True, timeout=30
            )
            # Simpler approach — just invoke ensure_webhook_subscriptions inline
            import importlib
            mod = importlib.import_module("onedrive_sync")
            graph = mod.GraphClient()
            try:
                await mod.ensure_webhook_subscriptions(p, graph)
            finally:
                await graph.close()
            
            return _json({"status": "ok", "message": "Webhook registration completed"})

        # GET — list existing webhooks
        rows = await p.fetch("""
            SELECT w.id, w.sync_target_id, w.subscription_id, w.resource,
                   w.change_type, w.expiration_at, w.created_at, w.renewed_at,
                   w.is_active, t.label as target_label, t.drive_owner
            FROM chat.onedrive_webhooks w
            LEFT JOIN chat.onedrive_sync_targets t ON t.id = w.sync_target_id
            ORDER BY w.created_at DESC
        """)
        return _json({"webhooks": [_row(r) for r in rows]})
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        import traceback; traceback.print_exc()
        return _json_error(f"Error: {e}", 500)


async def onedrive_webhook_delete(request: Request) -> JSONResponse:
    """Delete/deactivate a webhook subscription."""
    try:
        user = await _get_current_user(request)
        sub_id = request.path_params["sub_id"]
        p = await get_pool()
        await p.execute("""
            UPDATE chat.onedrive_webhooks SET is_active = false WHERE subscription_id = $1
        """, sub_id)
        return _json({"status": "ok"})
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        return _json_error(f"Error: {e}", 500)


async def onedrive_trigger_sync(request: Request) -> JSONResponse:
    """Trigger a OneDrive sync run."""
    try:
        user = await _get_current_user(request)
        body = await request.json() if request.method == "POST" else {}
        target_id = body.get("target_id")
        
        cmd = ["/opt/wdws/venv/bin/python3", "-u", "/opt/wdws/onedrive_sync.py"]
        if target_id:
            cmd.extend(["--target", str(target_id)])
        
        import subprocess
        proc = subprocess.Popen(
            cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True,
            env={**os.environ}
        )
        print(f"[ONEDRIVE] Triggered sync (PID={proc.pid}, target={target_id})")
        
        return _json({
            "status": "started",
            "pid": proc.pid,
            "target_id": target_id,
        })
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        return _json_error(f"Error: {e}", 500)


async def onedrive_toggle_target(request: Request) -> JSONResponse:
    """Toggle a sync target active/inactive."""
    try:
        user = await _get_current_user(request)
        target_id = int(request.path_params["target_id"])
        p = await get_pool()
        row = await p.fetchrow(
            "SELECT is_active FROM chat.onedrive_sync_targets WHERE id = $1", target_id)
        if not row:
            return _json_error("Target not found", 404)
        new_state = not row["is_active"]
        await p.execute(
            "UPDATE chat.onedrive_sync_targets SET is_active = $1 WHERE id = $2",
            new_state, target_id)
        return _json({"status": "ok", "is_active": new_state})
    except PermissionError as e:
        return _json_error(str(e), 401)
    except Exception as e:
        return _json_error(f"Error: {e}", 500)


routes = [
    Route("/health", health),
    Route("/quality-tiers", get_quality_tiers, methods=["GET"]),
    Route("/spaces", list_spaces, methods=["GET"]),
    Route("/spaces", create_space, methods=["POST"]),
    Route("/spaces/{space_id}", get_space, methods=["GET"]),
    Route("/spaces/{space_id}", update_space, methods=["PUT"]),
    Route("/spaces/{space_id}/links", list_space_links, methods=["GET"]),
    Route("/spaces/{space_id}/links", add_space_link, methods=["POST"]),
    Route("/spaces/{space_id}/conversations", list_conversations, methods=["GET"]),
    Route("/spaces/{space_id}/conversations", create_conversation, methods=["POST"]),
    
    # Space Knowledge Base
    Route("/spaces/{space_id}/documents", list_space_documents, methods=["GET"]),
    Route("/spaces/{space_id}/documents", add_space_document, methods=["POST"]),
    Route("/spaces/{space_id}/documents/{document_id}", get_space_document, methods=["GET"]),
    Route("/spaces/{space_id}/search", search_space_knowledge, methods=["GET"]),
    Route("/spaces/{space_id}/notes", list_space_notes, methods=["GET"]),
    Route("/spaces/{space_id}/notes", create_space_note, methods=["POST"]),
    Route("/spaces/{space_id}/stats", get_space_stats, methods=["GET"]),
    
    # OneDrive — Space linking
    Route("/spaces/{space_id}/onedrive", list_space_onedrive_links, methods=["GET"]),
    Route("/spaces/{space_id}/onedrive/link", link_onedrive_to_space, methods=["POST"]),
    Route("/spaces/{space_id}/onedrive/unlink", unlink_onedrive_from_space, methods=["DELETE"]),
    Route("/spaces/{space_id}/onedrive/search", search_space_onedrive, methods=["GET"]),
    
    # OneDrive — Global
    Route("/onedrive/targets", list_onedrive_targets, methods=["GET"]),
    Route("/onedrive/targets/{target_id}/toggle", onedrive_toggle_target, methods=["POST"]),
    Route("/onedrive/files", list_onedrive_files, methods=["GET"]),
    Route("/onedrive/search", search_onedrive, methods=["GET"]),
    Route("/onedrive/stats", get_onedrive_stats, methods=["GET"]),
    Route("/onedrive/sync", onedrive_trigger_sync, methods=["POST"]),
    Route("/onedrive/webhooks", onedrive_webhooks, methods=["GET", "POST"]),
    Route("/onedrive/webhooks/{sub_id}", onedrive_webhook_delete, methods=["DELETE"]),
    Route("/webhooks/graph", onedrive_webhook, methods=["GET", "POST"]),
    
    Route("/conversations/{conversation_id}", get_conversation, methods=["GET"]),
    Route("/conversations/{conversation_id}", update_conversation, methods=["PUT", "PATCH"]),
    Route("/conversations/{conversation_id}", delete_conversation, methods=["DELETE"]),
    Route("/conversations/{conversation_id}/messages", list_messages, methods=["GET"]),
    Route("/conversations/{conversation_id}/messages", create_message, methods=["POST"]),
    Route("/conversations/{conversation_id}/completions", create_completion, methods=["POST"]),
    Route("/conversations/{conversation_id}/completions/{request_id}/cancel", cancel_completion, methods=["POST"]),
    Route("/messages/{message_id}/edit", edit_message, methods=["PUT"]),
    Route("/messages/{message_id}/retry", retry_message, methods=["POST"]),
    Route("/messages/{message_id}/rate", rate_message, methods=["POST"]),
    Route("/messages/{message_id}/ratings", get_message_ratings, methods=["GET"]),
    Route("/assets", create_asset, methods=["POST"]),
    Route("/assets/{asset_id}/download", get_asset_download, methods=["GET"]),
    Route("/conversations/{conversation_id}/upload", upload_file, methods=["POST"]),
    Route("/images/generations", generate_image, methods=["POST"]),
    Route("/web_search", web_search, methods=["POST"]),
    Route("/memories", create_memory, methods=["POST"]),
    Route("/memories", list_memories, methods=["GET"]),
]

# ── Email Integration Routes ─────────────────────────────────
try:
    from email_api import get_email_routes
    routes.extend(get_email_routes())
except ImportError as e:
    print(f"[WARN] Email API not loaded: {e}")

middleware = [
    Middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_methods=["*"],
        allow_headers=["*"],
    )
]

app = Starlette(routes=routes, middleware=middleware)


def main():
    import uvicorn

    print("\n" + "=" * 60)
    print("  Athena Chat API")
    print(f"  http://{CHAT_HOST}:{CHAT_PORT}")
    print("=" * 60 + "\n")
    uvicorn.run(app, host=CHAT_HOST, port=CHAT_PORT)


if __name__ == "__main__":
    main()
